import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)

def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)
h2 = doc.add_heading('Mary Ajifowowe Oluwaseun', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Stoic Commuter')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201321, 2026 \u2022 5-Day Longitudinal Diary')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# ========== DIM 1 ==========
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Mary Ajifowowe Oluwaseun'],
    ['Location', 'Lagos \u2014 commutes from home to Ogba/Ojota/Alausa corridor via keke, bus, and bike'],
    ['Occupation', 'Employed worker \u2014 formal workplace with colleagues, receives orders and enquiries online'],
    ['Household', 'Young mother with at least one baby ("it was my baby\u2019s money"); mood-driven mornings'],
    ['Diary Period', '16\u201321 April 2026 (5 entries, some partial)'],
    ['Primary Network', 'MTN (main line) \u2014 "the best network so far," "fast network"'],
    ['Phone Behaviour', 'Does NOT use phone on waking most days (3/5) \u2014 but checks it as first act on 2/5'],
])

doc.add_paragraph(
    'Mary is the youngest-sounding respondent in the study. Her language is distinctly informal \u2014 '
    '"nothing tho," "nope," "nothing for now," emojis (\u270c\ufe0f, \u2639\ufe0f) \u2014 marking her as a '
    'different generational voice from the traders and shopkeepers who dominate the corpus. She '
    'works at a formal workplace where she has a desk ("arranged my table") and colleagues she '
    'buys food for. She is also a mother: on one entry, she explicitly refers to money spent as '
    '"my baby\u2019s money," revealing a layered financial life where her income is already earmarked '
    'for her child before she can spend on herself.'
)
doc.add_paragraph(
    'What makes Mary unique in this study is her combination of emotional minimalism and '
    'logistical complexity. She reports "nothing" to stress questions more often than any other '
    'respondent \u2014 not because her life is stress-free, but because she has developed a coping '
    'strategy of not naming her burdens. Her diary is full of quiet signals: she woke up "sick" '
    'one day and still went to work. She wanted to sleep at work but couldn\u2019t. She wanted to '
    'go to the market but work held her back. She rarely gets time to relax. Yet her consistent '
    'answer to "what stressed you?" is "nothing." Mary endures.'
)

# ========== DIM 2 ==========
add_h('Dimension 2: A Typical Day')

add_h3('Morning (05:00 \u2013 09:00)')
doc.add_paragraph(
    'Mary wakes between 5:00 and 6:00 AM. Her morning mood is a study in contrasts: she feels '
    '"so happy" on 3/5 days but woke up "sick" on 1/5, and still went to work. Her morning '
    'routine varies \u2014 on some days she prays and does house chores (2/5), on others she simply '
    'checks her phone (2/5). She uses her phone as a torch ("used the touch light") and for music '
    '("to play music") on one morning, revealing a sensory need \u2014 light and sound \u2014 to start her '
    'day. Her mood determines her morning: "It depends on my mood" is her consistent answer to '
    'what shapes her start, making her the most emotionally driven morning respondent.'
)
add_quote('I woke up late and just as I was about to leave home, Rain started')
doc.add_paragraph(
    'Her commute is the most detailed and multi-modal in the study. She navigates a complex '
    'transport chain: keke from her bus stop to Ogba, then another keke or bus to Ojota, Alausa, '
    'or First Gate, then sometimes walks the final stretch to her workplace. On good days, the '
    'journey is smooth \u2014 "Lagos is getting better." On one critical day, she took a bike instead '
    'of a keke "because I was late and wanted to avoid scolding," revealing workplace punctuality '
    'pressure and the real-time transport calculations she makes every morning.'
)
add_quote('I took bike from my bus stop to ogba, because I was late and wanted to avoid scolding')
add_quote('From my bus stop, I boarded a bus going to ogba, there was no traffic, Lagos is getting better')

add_h3('Afternoon (09:00 \u2013 16:00)')
doc.add_paragraph(
    'Mary\u2019s afternoons are defined by quiet endurance. Her language across multiple days reveals '
    'a profound monotony: "It\u2019s just been normal work hours," "It been just work and work tho," '
    '"Still at work, not really stressful, everything going smoothly," "It been just work and '
    'work, nothing exciting." These are not the words of someone who hates her job \u2014 they are '
    'the words of someone who has accepted that work will consume her daylight hours and finds '
    'no space for excitement within them.'
)
add_quote('It been just work and work, nothing exciting')
add_quote('I wanted to sleep but I couldn\u2019t not because am at work')
doc.add_paragraph(
    'Her phone helps her "confirm some information" and stay "online to receive any order or '
    'enquiries." She also uses her phone while commuting for bank transfers \u2014 "I was using it '
    'to do transfer" \u2014 revealing that she conducts financial business in transit, maximising '
    'every minute. On one day, she bought food for herself and a colleague during her commute, '
    'suggesting a generous, socially attentive personality despite her laconic diary style.'
)
add_quote('I was using it to do transfer')

add_h3('Evening (16:00 \u2013 22:00)')
doc.add_paragraph(
    'Mary\u2019s evenings are the sparsest in the study. On 2/5 days, she did "nothing" in the evening '
    'and had no relaxation time. On 1/5 days, she "went out with friends" \u2014 an activity she '
    'describes as happening "rarely." She browsed the internet on 1/5 evenings. She does not '
    'watch movies, does not use TikTok, does not game. Her evening is either empty or social, '
    'with nothing in between. She makes no dinner (unlike the traders), suggesting either someone '
    'else cooks or she eats outside. Her preparation for tomorrow is consistently "nothing" \u2014 '
    'she does not lay out clothes or plan. Mary lives in the present tense.'
)

# ========== DIM 3 ==========
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Daily', 'Get to work on time', 'Took a bike "because I was late and wanted to avoid scolding"'],
    ['Daily', 'Complete work duties', '"Work took most of my time today" \u2014 work IS the day'],
    ['Financial', 'Provide for her baby', '"it was my baby\u2019s money" \u2014 spending is child-first'],
    ['Personal', 'Sleep', '"I wanted to sleep but I couldn\u2019t" \u2014 a recurring unmet need'],
    ['Personal', 'Go to the market', '"I wanted to go to the market myself but I couldn\u2019t because of work\u2639\ufe0f"'],
    ['Social', 'Spend time with friends', 'Went out with friends (1/5 days) \u2014 "rarely" gets this time'],
])

# ========== DIM 4 ==========
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Frequency', 'Defining Quote'], [
    ['Work monotony / Boredom', '3/5 days', '"It been just work and work, nothing exciting"'],
    ['Sleep deprivation', '1/5 days (but implied)', '"I wanted to sleep but I couldn\u2019t not because am at work"'],
    ['Rain blocking commute', '1/5 days', '"Rain started" just as she was leaving home'],
    ['Lateness anxiety', '1/5 days', '"wanted to avoid scolding" \u2014 took bike to compensate'],
    ['No relaxation time', '2/5 days', '"Did you have time to relax? No"'],
    ['Sickness at work', '1/5 days', 'Woke up "sick" but still went to work'],
    ['Inability to do personal errands', '2/5 days', '"I wanted to go to the market," "laundry" \u2014 blocked by work'],
])
add_insight('Key Insight',
    'Mary\u2019s frustration profile is fundamentally different from every other respondent. She does '
    'not complain about power, network, or sales. Her pain is existential: the feeling that work '
    'absorbs her entire life, leaving no space for sleep, shopping, laundry, or social connection. '
    'She reports "nothing" to stress questions not because nothing stresses her, but because the '
    'stress is so ambient and constant that it has become invisible even to her. She went to work '
    'sick. She wanted to sleep but couldn\u2019t. She wanted to go to the market but couldn\u2019t. She '
    'rarely has time to go out with friends. Mary\u2019s pain point is not a specific failure \u2014 it '
    'is the total compression of her personal life by her work schedule.')

# ========== DIM 5 ==========
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network', 'MTN \u2014 "the best network so far," "fast network"'],
    ['Secondary network', 'None mentioned \u2014 likely single-SIM'],
    ['Multi-SIM', 'No'],
    ['Phone on waking', 'Mixed: No (3/5), Yes to check phone / play music (2/5)'],
    ['Phone as work tool', '"online to receive any order or enquiries," "confirm some information"'],
    ['Phone for transit', '"doing transfer" while commuting \u2014 mobile banking in motion'],
    ['Phone for leisure', 'Almost none: browsed internet (1/5), went out with friends (1/5)'],
    ['Phone as utility', 'Torch light (1/5), music player (1/5) \u2014 basic utility functions'],
    ['Phone charged', 'No (3/5 days) but no charging complaints \u2014 she doesn\u2019t flag it as stress'],
])
doc.add_paragraph(
    'Mary\u2019s phone relationship is the most understated in the study. She does not describe it '
    'as essential or critical \u2014 she simply uses it. She makes transfers while riding a keke. '
    'She stays "online to receive orders." She uses the torch to see in the dark. She plays '
    'music to start her day. But she never expresses anxiety about her phone dying, never '
    'complains about network, and answers "Nothing" when asked what tasks would be difficult '
    'without it (1/5 days). This is not because her phone is unimportant \u2014 it is because Mary '
    'does not dramatise. Her phone is like her commute: essential, routine, unremarkable in '
    'her own telling.'
)

# ========== DIM 6 ==========
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending frequency', '2/5 days \u2014 she spends infrequently'],
    ['Spending rationale', '"Because it was very important" / "money is to be spent and also be made"'],
    ['Child-first spending', '"it was my baby\u2019s money" \u2014 her earnings are allocated to her child first'],
    ['Generosity', 'Bought food for herself AND her colleague on one commute'],
    ['Payment method', 'Mobile transfers while commuting \u2014 digital-first'],
    ['Payment pain', 'None reported \u2014 "Nope, I didn\u2019t"'],
])
add_insight('Key Insight',
    'Mary\u2019s most revealing financial statement is "it was my baby\u2019s money." In five words, she '
    'reveals that her income is not hers \u2014 it belongs to her child first. When she spends, it '
    'must pass a double test: is it important, and does it take from her baby? This is not '
    'frugality \u2014 it is maternal sacrifice encoded into financial behaviour. Yet she also shows '
    'generosity: she bought food for a colleague without being asked. Mary\u2019s financial profile '
    'is one of quiet self-sacrifice \u2014 she gives to her child, gives to her colleague, and '
    'rarely spends on herself.')

# ========== DIM 7 ==========
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Primary method', 'Calls + WhatsApp \u2014 "it is affordable"'],
    ['Who she contacts', 'Customers (business) and Friends (personal) \u2014 distinct categories'],
    ['Call duration', '"Few minutes" \u2014 consistently brief'],
    ['Planned vs. spontaneous', 'Mix: planned business calls, planned social calls'],
    ['Fruitful?', 'Mixed: business call was "not fruitful" (1/2); social call was "yes" (1/2)'],
    ['Social time', '"went out with friends" + "just gisting" \u2014 rare but valued'],
])
doc.add_paragraph(
    'Mary\u2019s communication reveals an important split: her business calls are functional and '
    'sometimes unfruitful ("No" to "was it fruitful?"), while her social calls are warm and '
    'valued ("just gisting" with friends). She makes planned calls \u2014 unusual in a corpus '
    'dominated by spontaneous, reactive communication. This suggests she is deliberate about '
    'how she uses her airtime, perhaps because her budget is tight ("my baby\u2019s money"). '
    'Her social life is compressed but not absent: she went out with friends on one evening '
    'and described it as something she "rarely" gets to do.'
)
add_quote('Just gisting')

# ========== DIM 8 ==========
add_h('Dimension 8: Emotional Profile & Stress Map')

add_h3('Daily Emotional Arc')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Mostly happy (3/5), but sick (1/5)', 'She pushes through illness without complaint'],
    ['Morning commute', 'Calculated and alert', 'Real-time transport decisions (bike vs keke vs bus) based on lateness'],
    ['Work hours', 'Monotonous and accepting', '"work and work, nothing exciting" \u2014 endurance, not engagement'],
    ['Afternoon', 'Quietly frustrated', '"I wanted to sleep," "I wanted to go to market" \u2014 blocked desires'],
    ['Evening', 'Empty or social', 'Either "nothing" or rare, treasured outings with friends'],
])

add_h3('Stress Triggers (Ranked)')
doc.add_paragraph('1. Work monotony \u2014 the invisible weight of days that are "nothing exciting"')
doc.add_paragraph('2. Sleep deprivation \u2014 wanted to sleep at work but couldn\u2019t')
doc.add_paragraph('3. Lateness to work \u2014 fear of "scolding" drives expensive transport choices')
doc.add_paragraph('4. Rain \u2014 delays her departure and disrupts an already tight schedule')
doc.add_paragraph('5. Personal life compression \u2014 cannot do laundry, shopping, or socialising due to work')

add_h3('Resilience Pattern')
doc.add_paragraph(
    'Mary\u2019s resilience is stoic. She does not pray for strength (unlike the traders). She does not '
    'watch comedy to escape (unlike Adekunle). She does not gist with family to recharge (unlike '
    'Flora). Instead, she simply shows up. She went to work sick. She answered "nothing" to stress '
    'questions while clearly carrying stress. She finds joy in small moments: "the rain that fell" '
    'made her feel "productive, satisfied, or happy" on one day \u2014 a poetic response that suggests '
    'the rain was a welcome break from routine, a moment of beauty in an otherwise monotonous day. '
    'She also finds satisfaction in "Lagos is getting better" \u2014 an optimistic observation about '
    'traffic improvement that no other respondent made.'
)
add_quote('The rain that fell')
add_quote('Lagos is getting better')

# ========== DIM 9 ==========
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Commuter Music/Podcast Bundles', 'Mary plays music on her phone to start her day. A morning commuter bundle \u2014 '
     'data-free access to music or podcasts during peak commute hours (5\u20139 AM) \u2014 would serve '
     'her sensory need and create a daily brand ritual'],
    ['Late Worker Emergency Transport', 'She took a bike "to avoid scolding" when late. A partnership with ride-hailing '
     'services for discounted emergency rides during work commute hours would directly address '
     'her lateness anxiety and earn loyalty'],
    ['Baby & Working Mom Plans', '"It was my baby\u2019s money." A family-focused bundle that explicitly acknowledges '
     'working mothers \u2014 affordable data + baby-related content (health tips, parenting forums) \u2014 '
     'would resonate with her identity as a providing mother'],
    ['Mobile Banking Optimisation', 'She does transfers while riding a keke. Ultra-low-latency banking app access \u2014 '
     'or zero-rated banking data \u2014 would serve her in-transit financial behaviour'],
    ['Social Connection Bundles', 'She "rarely" gets time with friends. A "gisting bundle" \u2014 unlimited WhatsApp calling '
     'on weekend evenings \u2014 would help her maintain the social connections she clearly values '
     'but cannot afford in time or money'],
    ['Wellbeing Messaging', 'She went to work sick. She doesn\u2019t complain. She endures. Brand messaging that '
     'validates the silent strength of working mothers \u2014 "we see you" \u2014 would create emotional '
     'resonance that competitors cannot replicate'],
])

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('The rain still falling, am fully dressed, waiting for the rain to stop while I move but seems like it will last forever so am patiently waiting')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 Mary is fully dressed, ready to go, but trapped by the rain. She cannot leave. She '
    'cannot undress. She can only wait. This single image captures her entire experience: a woman '
    'who is always ready, always prepared, but constantly held back by forces she cannot control \u2014 '
    'the rain, the traffic, the work schedule, the baby\u2019s needs. She doesn\u2019t complain. She '
    'doesn\u2019t rage. She "patiently waits." The brand that sees this patience \u2014 and rewards it \u2014 '
    'will earn a customer who has never been seen before.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.save(f'{outdir}\\Persona_Mary_Ajifowowe_Oluwaseun.docx')
print("Persona saved: Mary Ajifowowe Oluwaseun (v2 deep-dive)")
