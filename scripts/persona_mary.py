import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Mary Ajifowowe Persona
add_heading_styled('Customer Persona: Mary Ajifowowe Oluwaseun', 0)
doc.add_paragraph('Persona Archetype: The Pragmatic Mother / "The Resilient Seller"')
doc.add_paragraph('Cluster Codes: A01, A05, B01, C01, C02, C06, D01, D04, E01, E02, F01, F08, H01, K03')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Female'],
    ['Occupation', 'Office/Sales Professional (Online orders & enquiries)'],
    ['Location Context', 'Lagos (Commutes to Alausa via bike and bus)'],
    ['Financial Context', 'Pragmatic; views money as a tool to be used. Manages household expenses, including "baby\'s money."'],
    ['Key Goal', 'Work/Business success; being online for customers while balancing motherhood.'],
    ['Household Role', 'Working Mother; cares for her baby while managing a demanding work schedule.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Mary’s day is a balance of professional duty and maternal responsibility. '
    'She wakes as early as 5:00 AM, starting with prayer (H01) and often checking her phone immediately '
    'to stay "online" for business enquiries (A01). Her mood determines her start, but "preparing for work" '
    'is always the most urgent task. Her commute is a tactical exercise—using bikes to avoid "scolding" '
    'when late and navigating the unpredictable Lagos weather. At work, she is focused on orders and '
    'business talk, while evenings are reserved for friends and family.'
)

doc.add_paragraph('**The Morning (5:00 AM - 9:00 AM)**')
doc.add_paragraph(
    'Mornings are for "getting set." She prays and prepares for work. '
    'On days she is late, she takes a "bike" instead of a bus to save time. '
    'Rain is a common disruptor, often forcing her to "patiently wait" for it to stop before moving. '
    'Her phone is often "not charged," yet it is the first thing she checks for work updates.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The core work period involves "receiving orders" and "confirming information." '
    'She is a supportive colleague, often "buying food" for herself and others. '
    'She relies on digital transfers (D04) to manage these small, essential spends. '
    'Despite feeling "sick" or "sleepy" at times, she remains "fully at work."'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Closing from work is often uneventful ("nothing exciting"). '
    'However, she finds joy in "meeting with family and friends" or "gisting" on WhatsApp. '
    'She rarely has time to relax deeply, often going out or spending time with her baby.'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Physical', 'Sickness / Tiredness (F08)', 'Causes a desire to "sleep" during work hours.'],
    ['Environmental', 'Rainfall (C06)', 'Delays the commute and complicates preparation.'],
    ['Social', 'Pressure of lateness', 'Forces more expensive transport choices (bikes) to avoid scolding.'],
    ['Infrastructure', 'Phone not charged', 'Adds a layer of preparation stress in the morning.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Mary is an MTN loyalist ("the best network so far"). She uses her phone primarily for '
    'work (WhatsApp/Calls) and financial transfers. She is "online" as a business strategy, '
    'ensuring she never misses a customer enquiry. Her tech usage is functional and professional.'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'Mary has a pragmatic relationship with money: "money is to be spent and also be made." '
    'She prioritizes "important" spends like transport and food. Her role as a mother is '
    'evident in her financial priorities, specifically mentioning "baby\'s money."'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Social and professional communication are intertwined. She gists with friends on WhatsApp '
    'but maintains a professional "business talk" profile for customers. Calls are her preferred '
    'method because of the "fast network."'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Mary is resilient and generally "happy." She finds silver linings even in the rain '
    '("the rain that fell" made her feel productive). Her stress is primarily driven by '
    'the logistics of "getting to work" and the physical toll of a long day.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **The Resilient Seller**: Staying online and available despite physical/weather barriers.')
doc.add_paragraph('2. **Maternal Pragmatism**: Balancing professional ambition with the needs of a child.')
doc.add_paragraph('3. **Lagos Optimism**: Finding "easy access" and "better days" in a challenging city.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **"Always Online" Data Bundles**: Low-cost, consistent connectivity for small-scale sellers.\n'
    '2. **Maternal Wellness/Efficiency**: Products or services that help working mothers manage time better.\n'
    '3. **Commute-Friendly Insurance**: Coverage or support for bike-based commuting in Lagos.'
)

# Save
doc_path = f'{outdir}\\Persona_Mary_Ajifowowe_Oluwaseun.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
