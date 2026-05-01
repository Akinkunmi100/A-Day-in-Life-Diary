import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Ehindero Sunday Persona
add_heading_styled('Customer Persona: Ehindero Sunday', 0)
doc.add_paragraph('Persona Archetype: The Optimistic Commuter / "Digital Optimist"')
doc.add_paragraph('Cluster Codes: A01, A05, B08, C01, C02, D04, E02, F01, F04, H01, H02, I05')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Male'],
    ['Occupation', 'Designer / Insurance Agent'],
    ['Location Context', 'Lagos (Commutes between home and office)'],
    ['Financial Context', 'Budget-conscious but digitally active; uses apps for transfers and work.'],
    ['Key Goal', 'Growth and professional development; being productive despite environmental constraints.'],
    ['Household Role', 'Likely young professional; interacts with colleagues and peer groups.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Ehindero’s day is anchored by ritual and discipline. He wakes at 6:00 AM, beginning with prayer and morning devotion (H01). '
    'His morning is a race against the Lagos commute—showering, ironing, and departing for the office (H02). '
    'His workday is split between insurance agency tasks and design work, heavily relying on his phone for communication and task management (A01). '
    'The commute back home is often a "battle" with traffic and heat (C01), leading to an evening of decompression through social browsing or rest (F06).'
)

doc.add_paragraph('**The Morning (6:00 AM - 9:00 AM)**')
doc.add_paragraph(
    'Mornings are focused and optimistic. "I felt happy and energetic," he notes on multiple days. '
    'The focus is on "morning devotion" and "getting set for office." '
    'The phone is used to check messages or "reach out to people" immediately upon waking.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'High productivity period. He is usually at the "office" or "in training." '
    'He uses his phone for "sending information to people," "checking designs," and "bank transfers." '
    'Social interaction (E02) with colleagues is a core part of this block.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'The commute home is the primary stressor. "Traffic was much," "the heat was too much." '
    'Once home, he relaxes by "browsing on Facebook" or "listening to music" before preparing for the next day (H03).'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Mobility', 'Heavy traffic and bad roads (C01)', 'Causes physical exhaustion and "stress peaks" (F02).'],
    ['Network', 'Intermittent slow network (B02)', 'Delayed bank transfers and interrupted social browsing.'],
    ['Financial', 'High cost of transportation (D01)', 'Increases the burden of the daily commute.'],
    ['Environment', 'Heat and noise during commute', 'Negative impact on emotional state during transition home.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Ehindero is a "Digital Optimist." Unlike other respondents who are frustrated by tech, '
    'he frequently reports "all was smooth" (B08). He uses Airtel for calling and MTN for evening data/streaming, '
    'demonstrating strategic Multi-SIM behaviour (B06). His phone is a tool for both "work productivity" and "evening decompression."'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'He is comfortable with digital finance (D04), frequently using bank apps for transfers. '
    'His spending is disciplined—focused on "fare to work" and "food." He shows "opportunity awareness" (I03), '
    'thinking about how to grow his insurance and design businesses.'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Social life is active and balanced. He communicates via WhatsApp and calls. '
    'Interactions are often "planned," particularly with clients and colleagues. '
    'He values "reaching out to people" to maintain professional and personal networks (E02, E04).'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Ehindero exhibits a high degree of "Resignation/Adaptation" (F05). '
    'While he notes stress (F02) from traffic, his overall tone remains positive. '
    'He experiences a "satisfaction peak" (F04) upon completing work tasks or reaching home safely.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **Survival Mobility**: The daily battle with Lagos infrastructure.')
doc.add_paragraph('2. **The Optimistic Professional**: Maintaining productivity despite constraints.')
doc.add_paragraph('3. **Digital Fluidity**: Seamlessly switching between apps for work and rest.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **Commuter-Optimised Content**: Short, engaging content for high-traffic commute periods.\n'
    '2. **Enterprise Growth Tools**: Low-cost professional tools for insurance/design agents.\n'
    '3. **Reliable Data Bundles**: Specifically targeting the 7:00 PM - 10:00 PM "decompression" window.'
)

# Save
doc_path = f'{outdir}\\Persona_Ehindero_Sunday.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
