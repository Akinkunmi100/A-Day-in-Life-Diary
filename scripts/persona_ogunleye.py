import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Ogunleye Odunayo Persona
add_heading_styled('Customer Persona: Ogunleye Odunayo', 0)
doc.add_paragraph('Persona Archetype: The Collaborative Professional / "The Supportive Anchor"')
doc.add_paragraph('Cluster Codes: A01, A05, B01, C02, D01, D04, E01, E02, F01, F04, H01, H02, J01, K02')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Female'],
    ['Occupation', 'Office Professional (Laptop-based work)'],
    ['Location Context', 'Lagos (Daily commute via bike and bus)'],
    ['Financial Context', 'Digitally savvy; relies on bank transfers for daily needs (like food) when cash is scarce.'],
    ['Key Goal', 'Professional excellence and communal support; ensuring her work and her colleagues’ work is completed.'],
    ['Household Role', 'Shares household responsibilities; active in morning prayers and chores.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Ogunleye’s day is defined by a blend of disciplined routine and workplace camaraderie. '
    'She wakes as early as 5:00 AM, anchoring her morning in prayer (H01) and household chores (K02) before a multi-modal commute. '
    'Her professional life is centered around office tasks using her laptop (A05), but she frequently goes beyond her own job description to assist colleagues. '
    'Evenings are a mix of closing from work, relaxing with digital entertainment (YouTube/movies), and preparing for the next day.'
)

doc.add_paragraph('**The Morning (5:00 AM - 9:00 AM)**')
doc.add_paragraph(
    'Mornings are productive and spiritual. "We pray; I go to have a bath." '
    'She uses her phone immediately, often as a "touch light" if there is no power (J01), '
    'and then to check messages from customers or friends. The commute involves a bike to the bus stop, '
    'then a public bus, often accompanied by music (G02) to maintain her "happy" mood.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The core of her day is "working at the office." She is laptop-dependent for her tasks. '
    'A recurring highlight is her collaborative nature: "helping one of my colleagues with her task because she is not feeling well." '
    'She uses her phone for "transfer payments for food," highlighting a reliance on digital finance (D04) due to a lack of physical cash.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Closing from work is the primary transition. Evenings are dedicated to "relaxation" and "entertainment." '
    'She engages with "Korea movies on YouTube" or "browses the internet." '
    'Preparation for the next day is also a priority: "I have chosen my clothes for work tomorrow" (H03).'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Infrastructure', 'No light / Power outages (J01)', 'Slows down preparation; causes "phone charging anxiety."'],
    ['Network', 'Bad/Poor network (B02)', 'Interrupted communication and "uncomfortable" feelings.'],
    ['Financial', 'Cash scarcity', 'Forces reliance on digital transfers for small purchases (food).'],
    ['Mobility', 'Traffic congestion', 'Causes stress peaks during the commute.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Ogunleye is a loyal GLO user ("that\'s the number people know me with"). '
    'She is a multi-device user, alternating between her laptop for work and her phone for "easy access" to her social network. '
    'She is a heavy consumer of video content (YouTube/Movies) for decompression.'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'Her spending is utilitarian: "transport" and "food." She is an early adopter of the "cashless" lifestyle by necessity, '
    'using her phone for "bank transfers" at restaurants when cash is unavailable. '
    'She prioritizes "getting to work" as her most necessary expense.'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Socially active via WhatsApp and in-person gisting. She maintains a balance between professional '
    'discussions and social interactions with friends and family. Communication is chosen for its "easy access."'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Ogunleye feels most "productive and satisfied" when she is helping others (colleagues or family). '
    'She starts her day with "happiness" but experiences dips when faced with "bad network" or "traffic." '
    'Her resilience is rooted in her routine and spiritual practice.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **The Supportive Anchor**: Success is measured by team/collective output.')
doc.add_paragraph('2. **Digital Necessity**: Tech is a lifeline for daily survival (payments/light).')
doc.add_paragraph('3. **Disciplined Ritual**: A day anchored by prayer and preparation.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **Micro-Payment Solutions**: Enhancing the reliability of small transfers for food/transport.\n'
    '2. **Colleague-Centric Workspaces**: Tools that facilitate easy task-sharing and collaboration.\n'
    '3. **Offline Entertainment Bundles**: Movie/video bundles that can be downloaded during "good network" periods for offline viewing.'
)

# Save
doc_path = f'{outdir}\\Persona_Ogunleye_Odunayo.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
