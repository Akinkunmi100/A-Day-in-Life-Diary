import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)

def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)
h2 = doc.add_heading('Adeola Olowolagba', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Besieged Trader')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201321, 2026 \u2022 6-Day Longitudinal Diary')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# ========== DIM 1 ==========
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Adeola Olowolagba'],
    ['Location', 'Lagos \u2014 walks to shop; treks to visit friends; commutes via bike and bus'],
    ['Occupation', 'Market trader / Shop owner \u2014 sells goods to both walk-in and online customers'],
    ['Household', 'Working mother: children in school + extended family (sister, church community)'],
    ['Diary Period', '16\u201321 April 2026 (6 entries across 6 consecutive days)'],
    ['Primary Network', 'Airtel (main line) + MTN (data/internet) \u2014 dual-SIM strategist'],
    ['Phone Behaviour', 'Uses phone immediately on waking (4/6 days) \u2014 to check time and customer messages'],
])

doc.add_paragraph(
    'Adeola is a Lagos market trader who exists at the intersection of three unrelenting pressures: '
    'the obligation to feed her children, the need to keep her business alive, and the daily battle '
    'against infrastructure failure. She runs a physical shop but also serves customers online via '
    'WhatsApp and calls, making her phone a direct revenue channel. She walks to her shop most days, '
    'takes her children to school, and returns home to cook, clean, and prepare for the next cycle. '
    'Her diary reveals a woman who thinks about money constantly \u2014 not in terms of accumulation, '
    'but in terms of survival arithmetic: every naira spent on food must be "replaced" through sales.'
)

# ========== DIM 2 ==========
add_h('Dimension 2: A Typical Day')

add_h3('Morning (05:30 \u2013 10:00)')
doc.add_paragraph(
    'Adeola wakes between 5:30 and 8:00 AM, though most days she is up by 6:00\u20136:30. Her mornings '
    'begin with prayer (5/6 days) and immediately transition into heavy domestic labour: house chores, '
    'cooking, and bathing. On school days, she takes her children to school before heading to the shop. '
    'Her morning mood splits evenly \u2014 she reports feeling "happy" (3/6) and "energetic" (1/6) on good '
    'days, but "relaxed" in a way that suggests resignation rather than calm on others. Her most urgent '
    'priority is always preparing for work, but the real bottleneck is the domestic load that must be '
    'cleared first.'
)
add_quote('After taking my children to school, I tidy up the house and went to shop')
doc.add_paragraph(
    'Traffic is a consistent antagonist. On one pivotal day, she reported being stuck: "I should be at '
    'work before 10am but can\u2019t because of traffic." This is not an inconvenience \u2014 it is lost revenue. '
    'Every hour she is not at the shop is an hour a customer might have called her number and gotten '
    'no answer.'
)
add_quote('I should be at work before 10am but can\u2019t because of traffic')

add_h3('Afternoon (10:00 \u2013 16:00)')
doc.add_paragraph(
    'The afternoon is the core business period. She is either "attending to customers" or, more often, '
    '"waiting for customers" \u2014 a phrase that recurs across her entries with quiet desperation. Her phone '
    'is the critical bridge: she uses it to "attend to customers online," to "confirm payments" via '
    'transfer alerts, and to "post my business" on social media. Without her phone, her business '
    'effectively goes dark.'
)
add_quote('My phone has been down and customer might have been calling my number')
add_quote('I use it to attend to customers online')
doc.add_paragraph(
    'Her biggest afternoon stressor is the collision between "no light" and "low battery." On multiple '
    'days, she wanted to charge her phone but couldn\u2019t due to power outages. This creates a cascading '
    'failure: dead phone \u2192 missed customer calls \u2192 lost sales \u2192 financial stress. She also reports '
    'that her phone "hangs" when she tries to use it, adding a device-quality layer to her frustration.'
)
add_quote('I want to charge my phone and no light')
add_quote('sometimes if I want to use my phone it use to hang and it stress me')

add_h3('Evening (16:00 \u2013 22:00)')
doc.add_paragraph(
    'Evenings are for the second domestic shift. She closes from work and returns to "cooking and washing." '
    'Her entertainment is modest: she watches "comedy" or "movies" on her phone (3/6 days) and browses '
    'WhatsApp. But she also uses her phone strategically in the evening \u2014 "posting my business" on social '
    'media, which she describes as both work and relaxation. She sleeps early when possible, and her '
    'preparation for the next day is focused on laundry, chores, and the persistent need to charge her phone.'
)
add_quote('I wanted to eat well. I couldn\u2019t because I needed to spend little')

# ========== DIM 3 ==========
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Daily', 'Make sales', '"To make sales" is her stated morning priority (3/6 days)'],
    ['Daily', 'Feed her children', '"It is a must I spend on food daily" \u2014 non-negotiable'],
    ['Business', 'Expand inventory', '"I want to add to my business but no money"'],
    ['Financial', 'Secure school fees and rent', '"School fees and house rent" are her top spending considerations'],
    ['Personal', 'Sleep and rest', '"I wanted to sleep but I couldn\u2019t because of home chores" (2/6 days)'],
    ['Spiritual', 'Study scripture', '"I want to study bible verse but am busy in shop"'],
])

# ========== DIM 4 ==========
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Frequency', 'Defining Quote'], [
    ['No light / Power outage', '4/6 days', '"no light" \u2014 prevents charging, causes missed customer calls'],
    ['Low sales / No sales', '5/6 days', '"when there\u2019s no sales at the shop"'],
    ['Phone hanging / Device quality', '1/6 days (but chronic)', '"if I want to use my phone it use to hang and it stress me"'],
    ['No money to invest in business', '1/6 days', '"I want to add to my business but no money"'],
    ['Traffic delays', '2/6 days', '"Traffic tied me down"'],
    ['Network issues', '3/6 days', '"When I want to browse and there is no network"'],
    ['No money for airtime', '1/6 days', '"I want to make call but I don\u2019t have money to buy card"'],
])
add_insight('Key Insight',
    'Adeola\u2019s frustrations form a vicious cycle: power outages kill her phone \u2192 dead phone means '
    'missed customer calls \u2192 missed calls mean lost sales \u2192 lost sales mean no money \u2192 no money '
    'means she can\u2019t buy airtime to call customers back. This is not a single pain point \u2014 it is a '
    'systemic failure loop where infrastructure, device quality, and financial scarcity compound each '
    'other. She is aware of this loop: "My phone has been down and customer might have been calling '
    'my number."')

# ========== DIM 5 ==========
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network (calls)', 'Airtel \u2014 "that\u2019s the number people know me with"'],
    ['Secondary network (data)', 'MTN \u2014 "fast network" for internet and browsing'],
    ['Multi-SIM', 'YES \u2014 strategic dual-SIM: Airtel for identity, MTN for speed'],
    ['Phone on waking', 'YES \u2014 checks time and customer messages (4/6 days)'],
    ['Phone as work tool', '"confirming payment," "attending to online customers," "posting my business"'],
    ['Phone for leisure', 'Movies, comedy, WhatsApp (3/6 days); browsing internet (3/6 days)'],
    ['Phone charged on waking', 'No (3/6 days) \u2014 power outages directly affect her business readiness'],
    ['Device quality', 'Phone hangs and causes stress \u2014 she needs a better device'],
])
doc.add_paragraph(
    'Adeola\u2019s dual-SIM strategy reveals a sophisticated understanding of network economics. She keeps '
    'Airtel as her "main line" because her customers know her by that number \u2014 switching would mean '
    'losing her business identity. But for data-heavy tasks like browsing and posting, she switches '
    'to MTN for its "fast network." This is not casual multi-SIM usage; it is a deliberate business '
    'decision. Her phone is not a luxury \u2014 it is her shop counter, her cash register (confirming '
    'transfers), and her marketing department (posting business).'
)

# ========== DIM 6 ==========
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending frequency', '5/6 days \u2014 she spends almost every day'],
    ['Primary spend', 'Food \u2014 "It is a must I spend on food daily"'],
    ['Secondary spend', 'Transport, airtime, data \u2014 "card and data because it was necessary"'],
    ['Spending logic', '"Before I spend on one thing I will sell for me to replace the money I spent"'],
    ['Savings anxiety', '"I could have saved the money" \u2014 every spend carries guilt'],
    ['Payment pain', 'Network issues while trying to browse/spend (3/6 days)'],
    ['Background obligations', '"School fees and house rent" \u2014 always weighing against daily spend'],
])
add_insight('Key Insight',
    'Adeola\u2019s financial behaviour is best described as "replacement economics." She does not budget in '
    'the traditional sense \u2014 instead, she operates on a simple rule: every naira spent must be earned '
    'back through the next sale. "Before I spend on one thing I will sell for me to replace the money '
    'I spent." This creates enormous pressure on her sales performance, because a slow day doesn\u2019t just '
    'mean less profit \u2014 it means she literally cannot afford to eat well. "I wanted to eat well. I '
    'couldn\u2019t because I needed to spend little."')

# ========== DIM 7 ==========
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Primary method', 'Calls + WhatsApp + In-person \u2014 uses "all methods to communicate"'],
    ['Who she contacts', 'Customers, family, church members, friends, sister'],
    ['Call duration', 'Ranges from 15 mins to 2 hours \u2014 highly variable'],
    ['Conversation style', 'Mix of planned (business) and spontaneous ("something came up")'],
    ['Topics', '"About market prices and availability," "life and business," "the sermon"'],
    ['WhatsApp as business tool', '"posting my business" and "attending to online customers"'],
])
doc.add_paragraph(
    'Adeola\u2019s communication is simultaneously professional and deeply personal. She discusses "market '
    'prices and availability" with customers in one breath and "the sermon" with church members in the '
    'next. Her longest conversation (2 hours) was with friends about "life and business" \u2014 revealing '
    'how, for her, personal relationships and business networks are inseparable. She reunited with '
    '"old classmates" at a party and they "talked about life" \u2014 these social moments are rare but '
    'deeply valued.'
)

# ========== DIM 8 ==========
add_h('Dimension 8: Emotional Profile & Stress Map')

add_h3('Daily Emotional Arc')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Generally positive: Happy (3/6), Energetic (1/6), Relaxed (2/6)', 'Starts each day with spiritual grounding and purpose'],
    ['Morning', 'Purposeful but rushed', 'Domestic chores, children to school, then race to shop'],
    ['Mid-day', 'Anxious', '"At shop waiting for customers" \u2014 the uncertainty of sales defines this period'],
    ['Afternoon', 'Frustrated or resigned', '"Low sales," "no light," "phone down" \u2014 compounding stressors'],
    ['Evening', 'Depleted but social', 'Closes shop, cooks, discusses with family. Comedy for escape'],
])

add_h3('Stress Triggers (Ranked)')
doc.add_paragraph('1. No sales / Poor sales \u2014 the existential threat to her family')
doc.add_paragraph('2. Power outages \u2014 the infrastructure failure that kills her phone and business')
doc.add_paragraph('3. "Not having money" \u2014 the emotional weight of financial scarcity')
doc.add_paragraph('4. Traffic \u2014 the time thief that steals business hours')
doc.add_paragraph('5. Phone hanging \u2014 the device that fails her when she needs it most')

add_h3('Resilience Pattern')
doc.add_paragraph(
    'Adeola\u2019s emotional state is directly indexed to her sales performance. She repeatedly states that '
    'she feels "productive, satisfied, or happy" specifically "when I get credited" or "when there\u2019s '
    'sales." Conversely, her stress is always financial: "no money," "poor sales," "no sales." Her '
    'resilience comes from two sources: (1) her spiritual practice \u2014 prayer is the first act of '
    'every single day, and she aspires to "study bible verse" even when busy, and (2) her social '
    'network \u2014 she "discussed" with family at church, gisted with friends, and visited her sister. '
    'These connections sustain her through the low-sales days.'
)

# ========== DIM 9 ==========
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Trader Power Packs', 'Adeola lost business on 4/6 days due to phone charging failures. A branded portable '
     'charging solution marketed to market women \u2014 affordable, rugged, high-capacity \u2014 would directly '
     'prevent her cascading failure loop (no light \u2192 dead phone \u2192 missed calls \u2192 lost sales)'],
    ['Sales-Linked Data Plans', 'She uses data to "attend to online customers" and "post my business." A data plan '
     'that rewards heavy business usage (e.g., free bonus data after WhatsApp business messages) '
     'would align with her replacement economics mindset'],
    ['Dual-SIM Loyalty Programs', 'She strategically splits Airtel (identity) and MTN (speed). A loyalty program that '
     'rewards her for keeping both active \u2014 rather than punishing her for not being single-network \u2014 '
     'would earn trust from this sophisticated segment'],
    ['Transaction Alert Reliability', 'Her most critical phone task is "confirming payment." Any failure in transaction '
     'SMS delivery directly costs her money. Ultra-reliable, instant transaction alerts would be '
     'the single most valued telecom feature for her segment'],
    ['Micro-Credit for Stock', '"I want to add to my business but no money." A micro-lending partnership that enables '
     'small traders to restock during slow periods would address her deepest aspiration'],
])

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('Before I spend on one thing I will sell for me to replace the money I spent')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 The survival arithmetic of a Lagos market woman, distilled into a single sentence. '
    'For Adeola, there is no savings buffer, no safety net. Every naira is a cycle: spend, sell, '
    'replace. When sales stop, the cycle breaks \u2014 and she cannot eat well, cannot buy airtime, '
    'cannot charge her phone. The brand that understands this cycle can earn her loyalty for life.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.save(f'{outdir}\\Persona_Adeola_Olowolagba.docx')
print("Persona saved: Adeola Olowolagba (v2 deep-dive)")
