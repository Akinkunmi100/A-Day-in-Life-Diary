import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Adekunle Adepetun Persona
add_heading_styled('Customer Persona: Adekunle Adepetun', 0)
doc.add_paragraph('Persona Archetype: The Vigilant Shopowner / "The Early Riser"')
doc.add_paragraph('Cluster Codes: A04, A05, B01, C06, E02, F06, G01, H01, J01, J02, K02, K05')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Male'],
    ['Occupation', 'Shop Owner / Trader (Located at place of residence)'],
    ['Location Context', 'Lagos (Minimal commute as shop is at home; interacts with local government/council)'],
    ['Financial Context', 'Business-focused spender; prioritizes shop permits and essential survival items (water).'],
    ['Key Goal', 'Protecting and growing his business while navigating local regulatory challenges and infrastructure gaps.'],
    ['Household Role', 'Likely the head of household; disciplined in morning rituals and shop management.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Adekunle’s day is characterized by discipline and vigilance. '
    'He wakes extremely early (4:30 AM - 5:00 AM) to a phone alarm, starting with '
    'prayer (H01) and house chores before opening his shop. His business is '
    'his primary focus, but he often faces external stressors like government '
    'council permit issues or environmental factors like rain. '
    'He is a "Digital Decompressor," relying on TikTok in the evenings to '
    'escape the stress of a long day "awaiting customers."'
)

doc.add_paragraph('**The Morning (4:30 AM - 9:00 AM)**')
doc.add_paragraph(
    'Mornings are "structured and prayerful." He uses his phone alarm to '
    'ensure an early start. After prayer and chores, he transitions immediately '
    'to "arranging goods to sell." He is "time cautious" and feels happy when '
    'the morning starts smoothly.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The afternoon is the core business period. He is often "sitting in the shop," '
    'waiting for customers. This period is prone to "unplanned stressors," '
    'such as interference from "government council workers" regarding permits. '
    'He uses his phone as a calculation tool and a crisis-management device, '
    'calling the "CDA Chairman" to mediate local disputes.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Closing the shop leads into a period of recovery. He "closes from work" '
    'and seeks relaxation through his phone, specifically "browsing TikTok." '
    'Preparation for the next day involves choosing work clothes and managing '
    'phone battery—often a struggle due to "faulty power banks" or power outages.'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Regulatory', 'Government council / Permit issues (A05)', 'Slows down business; prevents selling to customers.'],
    ['Infrastructure', 'No light / Faulty power bank (J01/J02)', 'Causes stress; forces the use of phone "torch" at work.'],
    ['Environmental', 'Rain (C06)', 'Affects the shop and delays business activities.'],
    ['Physical', 'Hunger during work hours (F08)', 'Causes a dip in energy and mood while awaiting customers.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Adekunle is an MTN loyalist. His phone is a "multi-tool": an alarm, a '
    'calculator, a flashlight, and a crisis-management tool. He is also a '
    'consumer of "Digital Leisure," using TikTok as his primary method to '
    '"relax and escape stress" in the evenings.'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'Finances are "business-first." He is willing to spend to "settle shop permit '
    'issues" with the council, viewing it as a necessary business expense. '
    'Daily spends are minimal, focused on "water" and essential survival.'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Tactical and local communication. He gists with "neighbors" and "church members," '
    'but his most critical calls are to the "CDA Chairman" or customers. '
    'He uses WhatsApp and calls for their "affordability" and "ease of use."'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Adekunle’s emotional state is resilient. He feels "so happy" in the morning '
    'but experiences "pressure" when his business is threatened. He is '
    '"satisfied" only when sales are made, reflecting a strong vocational identity.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **Vigilant Business Defense**: Actively managing local regulatory threats.')
doc.add_paragraph('2. **Digital Decompression**: TikTok as a necessary emotional escape.')
doc.add_paragraph('3. **Infrastructure Resilience**: Navigating work with faulty power banks and no light.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **Local Business Protection Services**: Legal or mediation support for small shopkeepers.\n'
    '2. **Rugged Power Solutions**: High-capacity, durable power banks for market-based workers.\n'
    '3. **TikTok Data Bundles**: Affordable, high-volume data specifically for video-based relaxation.'
)

# Save
doc_path = f'{outdir}\\Persona_Adekunle_Adepetun.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
