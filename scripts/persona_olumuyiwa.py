import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Olumuyiwa Folarin Persona
add_heading_styled('Customer Persona: Olumuyiwa Folarin', 0)
doc.add_paragraph('Persona Archetype: The Ritualistic Coordinator / "The Tired Taskmaster"')
doc.add_paragraph('Cluster Codes: A04, A05, B06, C01, C06, F08, H01, H02, H04, I01, J02, K05')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Male'],
    ['Occupation', 'Technical Coordinator / Audio Engineer (Handles "teams" and "sound")'],
    ['Location Context', 'Lagos (Mobile; moves between office, church, and event venues)'],
    ['Financial Context', 'Functional spender; focuses on transport and necessary business costs.'],
    ['Key Goal', 'Operational efficiency; managing teams and equipment to ensure "fruitful" results.'],
    ['Household Role', 'Likely head of household or primary provider; disciplined and routine-oriented.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Olumuyiwa’s life is a battle between disciplined ritual and chronic exhaustion. '
    'He wakes early (often 5:00 AM - 7:00 AM) to a rigid routine of prayer (H01) and exercise (H04), '
    'yet frequently reports feeling "tired" or "restless." His work is mobile and collaborative, '
    'requiring him to coordinate teams and manage technical sound equipment across various Lagos venues. '
    'He is "time cautious" and relies heavily on his phone to bridge the gap between locations.'
)

doc.add_paragraph('**The Morning (5:00 AM - 9:00 AM)**')
doc.add_paragraph(
    'Mornings are for "anchoring." Prayer and exercise are non-negotiable, even when feeling tired. '
    'He checks his phone immediately to "confirm the time" and look for "important messages." '
    'Preparation for work is the primary focus, often departing early to "arrive early" and set the tone for the day.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The core work period involves "team selection," "collating reports," and "technical sound checks." '
    'He is often in the field, using his phone to "put calls through to teams." '
    'He experiences "blocked goals" due to environmental factors like "rainfall" affecting sound check or '
    '"low batteries" on essential equipment.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Evenings are for "closing" and "recovery." He rarely uses his phone for internet-based entertainment, '
    'preferring to "sleep" or "play games" (offline). '
    'He is meticulous about the next day, always "choosing clothes for work tomorrow."'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Physical', 'Chronic Tiredness / Restlessness (F08)', 'Affects mood and energy levels throughout the day.'],
    ['Environmental', 'Rainfall / Weather (C06)', 'Disrupts technical sound work and live streaming.'],
    ['Infrastructure', 'Low battery / Power issues (J02)', 'Causes stress during critical "sound check" or work moments.'],
    ['Mobility', 'Traffic and long commutes', 'Adds to physical exhaustion.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Olumuyiwa uses his phone as a "coordination engine." He prefers an "alternative" network for calling '
    'because it offers "fast network." He views his phone as essential for "almost every task." '
    'However, he is a low consumer of internet-leisure, rarely browsing the internet for relaxation.'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'Finances are functional and focused on "transportation to venues." '
    'He is "satisfied" when sales are made, indicating a performance-linked income or business-owner mindset.'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Highly professional communication style. He uses calls and WhatsApp to coordinate teams and '
    'communicate with customers. Conversations are "fruitful" and "necessary," focused on business varieties.'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'His emotional state is closely tied to "operational success." '
    'He feels "satisfied and happy" when tasks are completed or "meeting with family/friends." '
    'He suffers from a "restlessness" that prevents the deep sleep he clearly craves.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **Ritualistic Professionalism**: Discipline (prayer/exercise) as a counter to chaos.')
doc.add_paragraph('2. **Chronic Exhaustion**: The physical toll of the "Lagos hustle."')
doc.add_paragraph('3. **Technical Coordination**: The phone as a command-and-control center.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **High-Performance Battery Solutions**: Portable power for field-based technical workers.\n'
    '2. **Sleep & Wellness Support**: Products or services targeting professional restlessness/insomnia.\n'
    '3. **Coordination SaaS**: Simple tools for field team management that work on low-bandwidth "fast" networks.'
)

# Save
doc_path = f'{outdir}\\Persona_Olumuyiwa_Folarin.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
