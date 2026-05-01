"""
Generate Word Document Personas for all 11 respondents.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
from collections import Counter

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
f_raw = outdir + r'\Merged A day in life diary.xlsx'
f_coded = outdir + r'\coded_corpus_full.xlsx'

df_raw = pd.read_excel(f_raw, engine='openpyxl', header=None)
headers = df_raw.iloc[1]
df = df_raw.iloc[2:].copy()
df.columns = range(len(df.columns))
df = df.reset_index(drop=True)
coded = pd.read_excel(f_coded, engine='openpyxl')

trivial = {'nothing','no','yes','nan','none','n/a','nope','','no,','yes,','no, nothing'}
fam_names = {'A':'Work-Driven','B':'Network-Focused','C':'Mobility-Stressed',
             'D':'Finance-Anxious','E':'Socially-Connected','F':'Emotionally-Complex',
             'G':'Digitally-Engaged','H':'Ritual-Anchored','I':'Aspiration-Driven',
             'J':'Power-Constrained','K':'Domestically-Centred'}

def get_top_responses(resp_df, col, n=3):
    vals = []
    for _, row in resp_df.iterrows():
        if col < len(row) and pd.notna(row[col]):
            v = str(row[col]).strip()
            if v.lower() not in trivial and len(v) > 2:
                vals.append(v)
    return vals[:n]

def get_all_responses(resp_df, cols):
    vals = []
    for _, row in resp_df.iterrows():
        for c in cols:
            if c < len(row) and pd.notna(row[c]):
                v = str(row[c]).strip()
                if v.lower() not in trivial and len(v) > 2:
                    vals.append(v)
    return vals

for resp_name in df[1].unique():
    print(f"Generating persona for: {resp_name}")
    resp_df = df[df[1] == resp_name].reset_index(drop=True)
    resp_coded = coded[coded['Respondent'] == resp_name]
    safe_name = resp_name.replace(' ', '_')
    entries = len(resp_df)
    
    # Extract key data
    dates = [str(row[2])[:10] for _, row in resp_df.iterrows() if pd.notna(row[2])]
    wake_times = get_top_responses(resp_df, 4, 7)
    first_acts = get_top_responses(resp_df, 5, 3)
    feelings = get_top_responses(resp_df, 7, 7)
    morning_det = get_top_responses(resp_df, 8, 3)
    phone_use = get_top_responses(resp_df, 10, 3)
    priorities = get_top_responses(resp_df, 15, 3)
    destinations = get_top_responses(resp_df, 18, 3)
    transport = get_top_responses(resp_df, 19, 3)
    journeys = get_top_responses(resp_df, 20, 3)
    activities = get_top_responses(resp_df, 39, 3)
    phone_help = get_top_responses(resp_df, 42, 3)
    spending = get_top_responses(resp_df, 45, 3)
    interactions = get_top_responses(resp_df, 49, 3)
    discussions = get_top_responses(resp_df, 50, 3)
    comm_method = get_top_responses(resp_df, 51, 3)
    call_network = get_top_responses(resp_df, 54, 7)
    net_reason = get_top_responses(resp_df, 55, 3)
    eve_activities = get_top_responses(resp_df, 70, 3)
    relax_acts = get_top_responses(resp_df, 73, 3)
    eve_phone = get_top_responses(resp_df, 75, 3)
    eve_network = get_top_responses(resp_df, 77, 7)
    content = get_top_responses(resp_df, 78, 3)
    stress_escape = get_top_responses(resp_df, 79, 3)
    
    stressors = get_all_responses(resp_df, [28,31,34,37,61,65,68])
    satisfiers = get_all_responses(resp_df, [60])
    blockers = get_all_responses(resp_df, [27,30,33,36,64,67])
    
    # Code analysis
    all_codes = []
    for _, row in resp_coded.iterrows():
        for c in ['Code_1','Code_2','Code_3','Code_4']:
            if pd.notna(row[c]) and row[c]:
                all_codes.append(row[c])
    fam_freq = Counter([c[0] for c in all_codes])
    code_freq = Counter(all_codes)
    top_fam = fam_freq.most_common(1)[0][0] if fam_freq else '?'
    archetype = fam_names.get(top_fam, 'Unclassified')
    
    # Network analysis
    call_net_counter = Counter([n.upper() for n in call_network])
    eve_net_counter = Counter([n.upper() for n in eve_network])
    primary_net = call_net_counter.most_common(1)[0][0] if call_net_counter else '?'
    evening_net = eve_net_counter.most_common(1)[0][0] if eve_net_counter else '?'
    is_multi_sim = primary_net != evening_net and primary_net != '?' and evening_net != '?'
    
    # ===== BUILD WORD DOCUMENT =====
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    
    def add_h(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
        return h
    
    def add_table(hdrs, rows_data):
        t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(hdrs):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for r_idx, row_data in enumerate(rows_data):
            for c_idx, val in enumerate(row_data):
                cell = t.rows[r_idx+1].cells[c_idx]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()
    
    def add_quote(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f'\u201c{text}\u201d')
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Title
    doc.add_paragraph()
    h = doc.add_heading('Customer Persona', level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
        run.font.size = Pt(32)
    h2 = doc.add_heading(resp_name, level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
        run.font.size = Pt(24)
    p = doc.add_paragraph()
    run = p.add_run(f'Persona Archetype: {archetype}')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    p = doc.add_paragraph()
    run = p.add_run(f'Diary Period: {dates[0] if dates else "?"} to {dates[-1] if dates else "?"} \u2022 {entries} entries')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()
    
    # Dim 1
    add_h('Dimension 1: Demographics & Life Context')
    loc_mentions = [r for r in get_all_responses(resp_df, [20,29,32,35,39]) if any(w in r.lower() for w in ['surulere','ojuwoye','ogba','yaba','festac','okota','mushin'])]
    location = loc_mentions[0] if loc_mentions else 'Lagos'
    occ_hints = get_top_responses(resp_df, 39, 3) + get_top_responses(resp_df, 40, 3)
    occupation = '; '.join(occ_hints[:2]) if occ_hints else 'Not specified'
    
    add_table(['Field', 'Detail'], [
        ['Name', resp_name],
        ['Location', f'Lagos \u2014 {location[:60]}'],
        ['Primary Activities', occupation[:80]],
        ['Diary Entries', f'{entries} entries ({dates[0] if dates else "?"} to {dates[-1] if dates else "?"})'],
        ['Primary Network', f'{primary_net} (calls) / {evening_net} (evening data)'],
    ])
    
    # Dim 2 - Day segments
    add_h('Dimension 2: A Typical Day')
    
    h3 = doc.add_heading('Morning (05:00 \u2013 12:00)', level=3)
    for run in h3.runs: run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    morning_text = f'Wakes at {wake_times[0] if wake_times else "?"} most mornings. '
    morning_text += f'First acts: {first_acts[0] if first_acts else "?"}. '
    morning_text += f'Morning feeling: {", ".join(set([f.lower() for f in feelings[:4]]))}. '
    morning_text += f'Morning priority: {priorities[0] if priorities else "?"}. '
    if destinations and destinations[0].lower() not in trivial:
        morning_text += f'Leaves home for {destinations[0].lower()}. '
    if transport:
        morning_text += f'Transport: {transport[0].lower()}. '
    if journeys:
        morning_text += f'Journey: {journeys[0]}'
    doc.add_paragraph(morning_text)
    if journeys:
        add_quote(journeys[0])
    
    h3 = doc.add_heading('Afternoon (12:00 \u2013 18:00)', level=3)
    for run in h3.runs: run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    aftn_text = f'Main activities: {activities[0] if activities else "?"}. '
    if phone_help:
        aftn_text += f'Phone helps with: {phone_help[0]}. '
    if spending:
        aftn_text += f'Spending: {spending[0]}. '
    if interactions:
        aftn_text += f'Interacts with: {interactions[0]}. '
    if discussions:
        aftn_text += f'Discusses: {discussions[0][:80]}'
    doc.add_paragraph(aftn_text)
    if phone_help:
        add_quote(phone_help[0])
    
    h3 = doc.add_heading('Evening (18:00 \u2013 22:00)', level=3)
    for run in h3.runs: run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    eve_text = f'Evening activity: {eve_activities[0] if eve_activities else "?"}. '
    if relax_acts:
        eve_text += f'Relaxation: {relax_acts[0]}. '
    if eve_phone:
        eve_text += f'Phone for: {eve_phone[0]}. '
    if content:
        eve_text += f'Content: {content[0]}. '
    if stress_escape:
        eve_text += f'Effect: {stress_escape[0]}.'
    doc.add_paragraph(eve_text)
    
    # Dim 4 - Pain Points
    add_h('Dimension 3: Frustrations & Pain Points')
    stress_counter = Counter()
    for s in stressors:
        for w in s.lower().replace(',',';').split(';'):
            w = w.strip()
            if w and w not in trivial and len(w) > 2:
                stress_counter[w] += 1
    pain_rows = [[s.capitalize(), str(c)] for s, c in stress_counter.most_common(6)]
    if pain_rows:
        add_table(['Stressor', 'Mentions'], pain_rows)
    if stressors:
        add_quote(stressors[0])
    
    # Dim 5 - Network
    add_h('Dimension 4: Phone & Network Relationship')
    net_rows = [
        ['Primary network (calls)', f'{primary_net} \u2014 {net_reason[0] if net_reason else "?"}'],
        ['Evening network (data)', evening_net],
        ['Multi-SIM', 'Yes' if is_multi_sim else 'No'],
        ['Phone first act on waking', 'Yes' if phone_use else 'No'],
    ]
    if phone_use:
        net_rows.append(['Morning phone use', phone_use[0]])
    if phone_help:
        net_rows.append(['Phone at work', phone_help[0][:60]])
    add_table(['Aspect', 'Detail'], net_rows)
    
    # Dim 6 - Financial
    add_h('Dimension 5: Financial Behaviour')
    if spending:
        doc.add_paragraph(f'Spending pattern: {spending[0]}')
    pain_payment = get_all_responses(resp_df, [48])
    if pain_payment:
        doc.add_paragraph(f'Payment pain: {pain_payment[0]}')
        add_quote(pain_payment[0])
    
    # Dim 7 - Communication
    add_h('Dimension 6: Communication Style')
    comm_rows = []
    if comm_method:
        comm_rows.append(['Primary method', comm_method[0]])
    if interactions:
        comm_rows.append(['Who', interactions[0]])
    if discussions:
        comm_rows.append(['Topics', discussions[0][:80]])
    call_dur = get_top_responses(resp_df, 58, 3)
    if call_dur:
        comm_rows.append(['Call duration', call_dur[0]])
    if comm_rows:
        add_table(['Aspect', 'Detail'], comm_rows)
    
    # Dim 8 - Emotional
    add_h('Dimension 7: Emotional Profile')
    feel_counter = Counter([f.lower() for f in feelings])
    feel_rows = [[f.capitalize(), str(c)] for f, c in feel_counter.most_common(5)]
    if feel_rows:
        add_table(['Morning Feeling', 'Days'], feel_rows)
    if satisfiers:
        doc.add_paragraph(f'What brings satisfaction: {satisfiers[0]}')
    
    # Dim 9 - Code Profile
    add_h('Dimension 8: Thematic Code Profile')
    fam_full = {'A':'Work & Productivity','B':'Network Experience','C':'Mobility & Movement',
                'D':'Financial Behaviour','E':'Social & Relational','F':'Emotional Experience',
                'G':'Digital Leisure','H':'Routine & Spiritual','I':'Aspirations & Unmet Needs',
                'J':'Power & Infrastructure','K':'Domestic Labour'}
    fam_rows = [[f'{f} \u2014 {fam_full.get(f,"")}', str(c), f'{c/max(sum(fam_freq.values()),1)*100:.0f}%'] 
                 for f, c in fam_freq.most_common()]
    add_table(['Family', 'Count', '%'], fam_rows)
    
    top_code_rows = [[c, str(n)] for c, n in code_freq.most_common(10)]
    add_table(['Top Code', 'Count'], top_code_rows)
    
    # Save
    doc_path = f'{outdir}\\Persona_{safe_name}.docx'
    try:
        doc.save(doc_path)
        print(f"  Saved: {doc_path}")
    except PermissionError:
        alt_path = f'{outdir}\\Persona_{safe_name}_v2.docx'
        doc.save(alt_path)
        print(f"  Original locked, saved as: {alt_path}")

print(f"\nAll {len(df[1].unique())} persona documents generated!")
