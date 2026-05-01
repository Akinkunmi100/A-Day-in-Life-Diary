import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Atolagbe Flora Yemi Persona
add_heading_styled('Customer Persona: Atolagbe Flora Yemi', 0)
doc.add_paragraph('Persona Archetype: The Environmental Survivor / "The Family-Centric Shopkeeper"')
doc.add_paragraph('Cluster Codes: A01, B01, C01, C06, D01, D02, E01, E02, F03, H01, J01, J02, K01, K02')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Female'],
    ['Occupation', 'Shop Owner / Trader (Physical retail)'],
    ['Location Context', 'Lagos (Walks to shop with her children; commute affected by road quality)'],
    ['Financial Context', 'Highly protective of earnings; avoids unnecessary spend when sales are poor.'],
    ['Key Goal', 'Ensuring the survival and comfort of her family through her retail business, despite infrastructure challenges.'],
    ['Household Role', 'Primary Caregiver and Domestic Lead; responsible for cooking, chores, and raising children.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Flora’s day is a tapestry of domestic labor and retail grit. '
    'She wakes between 6:30 AM and 8:00 AM, beginning her day with prayer (H01) '
    'and immediately transitioning into a heavy load of house chores and cooking. '
    'Her workday is spent at her shop, which she reaches by walking with her children. '
    'Her business is acutely sensitive to the environment—rain and bad roads directly '
    'impact her patronage and energy levels. Evenings are dedicated to family gisting, '
    'preparing dinner, and managing the chronic anxiety of phone charging.'
)

doc.add_paragraph('**The Morning (6:30 AM - 10:00 AM)**')
doc.add_paragraph(
    'Mornings are for "family and preparation." She wakes her family, prays, cooks, and tidies. '
    'The commute is a shared family journey: "I walked to shop with my kids." '
    'This journey is often "stressful" due to "muddy and rough" roads following rainfall (C06). '
    'Power supply (J01) often determines how her morning starts and how she prepares.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The core work period is spent "patiently attending to customers." '
    'Her phone is her primary business verification tool; she uses it to "carefully verify transaction alerts" '
    'from customers paying via transfer. Poor sales (D02) due to weather or road conditions '
    'is her biggest afternoon stressor, leading to "uncomfortable" and "pressured" feelings.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Closing from work leads back into domesticity. "Made dinner" is the primary evening task. '
    'Relaxation is social, consisting of "communicating with family or friends" in person. '
    'Preparation for the next day almost always centers on "charging my phone" (J02) '
    'to ensure business continuity.'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Environmental', 'Rainfall / Muddy roads (C06)', 'Disrupts commute and severely reduces shop patronage.'],
    ['Financial', 'Poor sales / No sales (D02)', 'Causes emotional stress and restricts household spending.'],
    ['Infrastructure', 'No light / Charging issues (J01/J02)', 'Disrupts the ability to confirm business transactions.'],
    ['Physical', 'Weakness / Weak body (F08)', 'Leads to "tiredness" and inability to perform morning prayers.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Flora is an MTN user, choosing it for "affordable bundles" and "fast network." '
    'Her phone usage is strictly utilitarian: "confirming transaction alerts" and '
    '"checking up on her husband/mother/siblings." She rarely uses her phone for '
    'entertainment, viewing it instead as a critical professional link.'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'Flora is a "Protective Spender." She prefers to eat food available at home rather than '
    'spend the "little money" made on a slow sales day. Her spending is focused on '
    '"dinner items" and necessary "transportation" when walking is not feasible.'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Deeply relational communication. She prioritizes gisting with her family '
    'and "checking up" on her mother and siblings via calls. Business communication '
    'is functional, focused on verifying "customer calls" and "transaction alerts."'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Flora experiences a high degree of "Environmental Resignation." '
    'She accepts the "rough roads" and "no light" as part of her daily struggle. '
    'She feels "productive and satisfied" when she can meet with family '
    'or when sales are healthy, but is prone to "weakness" when tired.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **Environmental Sensitivity**: Livelihood and mobility dictated by the weather.')
doc.add_paragraph('2. **Intertwined Domesticity**: Business labor that includes childcare (walking kids to shop).')
doc.add_paragraph('3. **Transaction Anxiety**: The stress of verifying digital payments on low-battery devices.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **Weather-Resilient Retail Support**: Services that help traders reach customers when the "rough roads" block them.\n'
    '2. **Transaction Verification SMS**: Ultra-reliable, low-power alerts for business owners to confirm payments.\n'
    '3. **Family-Centric Loyalty Programs**: Bundles that offer affordable calling specifically to family members (Mother/Husband).'
)

# Save
doc_path = f'{outdir}\\Persona_Atolagbe_Flora_Yemi.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
