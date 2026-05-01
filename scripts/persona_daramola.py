import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Daramola Solomon Persona
add_heading_styled('Customer Persona: Daramola Solomon', 0)
doc.add_paragraph('Persona Archetype: The Mobile Broker / "The Data-Driven Marketer"')
doc.add_paragraph('Note: This persona is synthesized from a limited dataset (1 day entry).')
doc.add_paragraph('Cluster Codes: A05, B01, C01, C06, D01, F08, G01, H01, K05')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Male'],
    ['Occupation', 'Insurance Professional / Marketer (Brokers & Underwriters)'],
    ['Location Context', 'Lagos (Mobile; moves between clients and underwriters via public transport)'],
    ['Financial Context', 'Data-focused; considers internet connectivity a "necessity" for service delivery.'],
    ['Key Goal', 'Ensuring policy renewals and effective service delivery through constant client follow-ups.'],
    ['Household Role', 'Likely an independent professional; handles own chores and meals.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Daramola’s day is built around professional appointments and client service. '
    'He wakes at 7:00 AM, anchoring his morning in a "cup of coffee" and a bath. '
    'His workday is mobile, involving "marketing runs" to insurance brokers and '
    'underwriters across Lagos. He relies heavily on his phone for coordination '
    'and views data as a critical business input. His day is characterized by '
    'the pressure of managing "new tasks" that interrupt his planned schedule.'
)

doc.add_paragraph('**The Morning (7:00 AM - 10:00 AM)**')
doc.add_paragraph(
    'Mornings are "functional." He starts with coffee, checks his phone, and '
    'prepares for work. He is a commuter, using public transportation for a '
    'short 10-minute journey to his primary work location. Rain is a primary '
    'stressor that can disrupt his early momentum.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The core work period involves "marketing runs" and "following up on '
    'outstanding tasks" like policy renewals. He is client-facing and '
    'appointment-driven, which makes "keeping up with appointment times" '
    'his most effort-intensive activity. He uses his phone while moving '
    'to check information or call contacts.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Evenings are for "closing" and decompression. He handles domestic '
    'chores like "making dinner" and "laundry." Relaxation is digital, '
    'consisting of "watching videos online" or "movies" on his phone.'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Workload', 'Task switching / Interruptions (A05)', 'Causes stress when new tasks derail outstanding work.'],
    ['Financial', 'Cash necessity for small payments', 'Frustration with the lack of digital payment options for certain services.'],
    ['Environmental', 'Rain and Traffic', 'Disrupts marketing runs and appointment timings.'],
    ['Physical', 'Tiredness upon waking (F08)', 'Affects initial productivity levels.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Daramola is an MTN loyalist. He is a "High-Value Data User," viewing '
    'data as a non-negotiable professional tool. He uses WhatsApp and '
    'calls extensively for client management because it is "effective and '
    'less time consuming" than physical visitation alone.'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'His spending is driven by "professional necessity." Data and transportation '
    'are prioritized to ensure "effective service delivery." He prefers digital '
    'finance and finds cash-only transactions to be a "pain point."'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Professional and proactive. He spends his day calling and chatting '
    'with "customers and Brokers" to discuss "work and policy renewals." '
    'Conversations are typically short (average 3 minutes) but "fruitful."'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Daramola is performance-oriented. He feels "satisfied and happy" '
    'when sales are made. His stress is situational, driven by "transportation '
    'delays" or "interruptions" to his workflow.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **The Data-Driven Marketer**: Digital connectivity as a professional prerequisite.')
doc.add_paragraph('2. **The Appointment-Driven Hustle**: Managing client expectations across locations.')
doc.add_paragraph('3. **Digital Finance Preference**: Resistance to the "cash necessity" of daily Lagos life.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **Professional Data Bundles**: Connectivity packages optimized for high-volume WhatsApp/Calling.\n'
    '2. **Digital Insurance Platforms**: Tools that allow brokers to renew policies without physical visitation.\n'
    '3. **Cash-to-Digital Micro-Payments**: Solutions that eliminate the need for cash during "marketing runs."'
)

# Save
doc_path = f'{outdir}\\Persona_Daramola_Solomon.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
