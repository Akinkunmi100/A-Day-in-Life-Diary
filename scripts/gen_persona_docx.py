import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return t

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)
    return p

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)

h2 = doc.add_heading('Adekoya Adesola', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)

p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Hustling Provider')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201322, 2026 \u2022 7-Day Longitudinal Diary')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ========== DIMENSION 1 ==========
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table(
    ['Field', 'Detail'],
    [
        ['Name', 'Adekoya Adesola'],
        ['Location', 'Lagos \u2014 operates around Ojuwoye, Surulere, Festac corridors'],
        ['Occupation', 'Self-employed trader / field worker (interviewing, photography, product advertising)'],
        ['Household', 'Lives with family; supports parents financially'],
        ['Diary Period', '16\u201322 April 2026 (7 consecutive days)'],
        ['Connection Type', 'Online (confirmed all 7 days)'],
    ]
)

# ========== DIMENSION 2 ==========
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Adesola wakes at 7:00 AM most mornings. She begins every day with prayer and house chores \u2014 '
    'this devotional-domestic sequence is unbroken across all seven diary entries. Her phone becomes active '
    'immediately. She checks messages, contacts customers, and leaves for work via public transport. Her '
    'days are structured around a cycle of work, errands for family, and social connection \u2014 punctuated '
    'by systemic friction from transport, network failures, and weather.'
)

# Morning subsection
h = doc.add_heading('Morning (05:00 \u2013 12:00)', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
doc.add_paragraph(
    'Every morning follows the same ritual: prayer, house chores, bathing, then preparing for work. '
    'This devotional-domestic sequence (H01 + K02 + K05) is unbroken across all seven days. Her phone '
    'is the first tool she reaches for \u2014 checking messages and contacting customers before she even '
    'leaves the house. She uses public transportation daily (bus, tricycle, BRT), and describes the '
    'commute as consistently stressful: bad roads, noise, queueing for buses, and delayed change from '
    'drivers. She always uses her phone during transit \u2014 checking information, making calls, or '
    'playing games. By mid-morning, blocked goals start accumulating: items not available at stores, '
    'people unreachable, weather disrupting plans.'
)
add_quote('It was a little bit stressful due to noise and bad road')
add_quote('A lot of thng slow me down in the area of transportation and network')

# Afternoon subsection
h = doc.add_heading('Afternoon (12:00 \u2013 18:00)', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
doc.add_paragraph(
    'Afternoons are her most productive but also most frustrating period. She juggles work tasks '
    '(WhatsApp advertising, photography, customer contact) with family errands (buying food for '
    'parents at Ojuwoye market, paying NEPA bills, picking up packages). Financial transactions are '
    'a recurring pain point \u2014 her GTB banking app fails due to network in 4 out of 7 entries, '
    'forcing her to fall back to cash. She spends money every day, always on necessities (food, '
    'transport, data), and always after considering her budget first. Her longest interactions happen '
    'in the afternoon: planned 3.5-hour voice calls with family about career opportunities, '
    'community development, and textile production.'
)
add_quote('My bank app had issues due to bad network')
add_quote('We discussed about making money and creating opportunities for young people in the community')

# Evening subsection
h = doc.add_heading('Evening (18:00 \u2013 22:00)', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
doc.add_paragraph(
    'Evenings are Adesola\u2019s recovery zone. She prepares clothes for the next day (every single day \u2014 '
    'H03), then retreats into digital leisure. She plays mobile games, watches movies, listens to music, '
    'and browses food content. She explicitly describes this as \u201cEscape Stress\u201d (5 of 7 entries). '
    'Notably, she switches from GLO to MTN for evening internet/streaming on some days \u2014 classic '
    'multi-SIM behaviour driven by data quality for entertainment. Her evenings end with family '
    'discussions or playing games with friends, a deliberate social decompression before sleep.'
)
add_quote('Escape Stress')
add_quote('I have chosen my clothes for work tomorrow')

# ========== DIMENSION 3 ==========
add_heading_styled('Dimension 3: Goals & Motivations')
add_table(
    ['Type', 'Goal', 'Evidence'],
    [
        ['Daily', 'Get to work on time', 'Coded as most urgent priority in 6/7 entries'],
        ['Daily', 'Complete errands for parents', '\u201cTo buy food stuff for my parent\u201d'],
        ['Business', 'Advertise products and attract customers', '\u201cSending of the product to all my WhatsApp contacts as a method of Advertising\u201d'],
        ['Aspirational', 'Create opportunities for community', '\u201cWe discussed about making money and creating opportunities for young people\u201d \u2014 3.5 hour planned call'],
        ['Aspirational', 'Career in fabrics/textiles', '\u201cDiscussed about production of fabrics and textile products\u201d'],
    ]
)

# ========== DIMENSION 4 ==========
add_heading_styled('Dimension 4: Frustrations & Pain Points')
add_table(
    ['Pain Point', 'Frequency', 'Defining Quote'],
    [
        ['Transportation stress', '7/7 days', '\u201cA lot of thng slow me down in the area of transportation and network\u201d'],
        ['Bank app failure (GTB)', '4/7 days', '\u201cMy bank app had issues due to bad network\u201d'],
        ['Bad weather blocking movement', '5/7 days', '\u201cI wanted to go to the hospital, but the rain could not allow me\u201d'],
        ['Unreachable contacts', '5/7 days', '\u201cI wanted to call a friend for assistance but his phone was not available\u201d'],
        ['Service gaps', '6/7 days', '\u201cI wanted to meet the customer care at the Airtel office but she was on leave\u201d'],
        ['Ride app failure', '2/7 days', '\u201cI wanted to order ride on my app but it was not going through\u201d'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('Key Insight: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
run = p.add_run('The I01 (Unmet Task) code appeared 27 times across 7 days. Adesola is constantly trying '
    'to do things and being blocked \u2014 by network, weather, transport, or people not being available. '
    'This is the most defining behavioural pattern: persistent effort meeting systemic failure.')

# ========== DIMENSION 5 ==========
add_heading_styled('Dimension 5: Phone & Network Relationship')
add_table(
    ['Aspect', 'Detail'],
    [
        ['Primary network', 'GLO \u2014 \u201cbundle is affordable\u201d, \u201cfast network\u201d'],
        ['Secondary network', 'MTN \u2014 used for evening internet/streaming'],
        ['Multi-SIM behaviour', 'Confirmed \u2014 GLO for daytime calls; MTN for entertainment'],
        ['Phone as first act', 'Checks messages immediately on waking (7/7 days)'],
        ['Phone at work', 'WhatsApp commerce, photography, product info, customer contact'],
        ['Phone while commuting', 'Always used during transit: information, calls, games'],
        ['Phone for evening', 'Movies, games, music, food content'],
        ['Network pain', 'Recurring GTB app failures; network blocks financial transactions'],
    ]
)

# ========== DIMENSION 6 ==========
add_heading_styled('Dimension 6: Financial Behaviour')
add_table(
    ['Aspect', 'Detail'],
    [
        ['Spending pattern', 'Spends money every day (7/7) \u2014 food, transport, data'],
        ['Budget discipline', '\u201cFirst and foremost I had to consider my budget\u201d'],
        ['Digital payment', 'Prefers bank app but frequently blocked by network'],
        ['Cash as workaround', '\u201cMy app couldn\u2019t allow me make transfer so I gave them cash\u201d'],
        ['Financial stress', 'Slow sales; \u201cno money\u201d as stressor'],
        ['Provider role', 'Buys for parents, pays NEPA bills, supports siblings'],
    ]
)

# ========== DIMENSION 7 ==========
add_heading_styled('Dimension 7: Communication Style')
add_table(
    ['Aspect', 'Detail'],
    [
        ['Primary channel', 'Voice calls (7/7 entries)'],
        ['Secondary channel', 'WhatsApp \u2014 for business advertising'],
        ['Who she contacts', 'Family (dominant), friends, customers, community'],
        ['Call duration', '3 hours 30 minutes (reported in 3 entries)'],
        ['Planned vs spontaneous', 'Primarily planned conversations'],
        ['Topics', 'Money-making, career opportunities, community development, social issues'],
    ]
)

# ========== DIMENSION 8 ==========
add_heading_styled('Dimension 8: Emotional Profile & Stress Map')
add_table(
    ['Time Block', 'Typical Emotion', 'Evidence'],
    [
        ['Wake-up', 'Mixed \u2014 Sick (3 days) or Energetic (4 days)', 'Alternates between unwellness and energy'],
        ['Morning commute', 'Stressed', 'Transportation is most consistent stressor'],
        ['Mid-morning', 'Frustrated', 'Blocked goals pile up'],
        ['Afternoon', 'Anxious', '\u201cAnxiety\u201d reported in 4/7 entries'],
        ['Evening', 'Decompressed', '\u201cEscape Stress\u201d \u2014 5/7 entries'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('Resilience Pattern: ')
run.bold = True
p.add_run('Adesola doesn\u2019t complain \u2014 she adapts. When the app fails, she pays cash. When the ride '
    'app breaks, she queues for the bus. When the rain blocks her, she changes plans. Her resilience is '
    'not passive acceptance \u2014 it is active problem-solving in real time.')

# ========== DIMENSION 9 ==========
add_heading_styled('Dimension 9: Opportunities for the Brand')
add_table(
    ['Opportunity', 'Actionable Insight'],
    [
        ['Fix the payment moment', 'A low-bandwidth payment solution would earn deep loyalty'],
        ['Evening data bundles', 'She switches to MTN for streaming \u2014 affordable GLO evening bundle prevents SIM switching'],
        ['Ride-hail reliability', 'Network optimisation for transport corridors reduces daily pain'],
        ['Community connector programme', '3.5-hour calls about youth opportunities = natural influencer'],
        ['WhatsApp Business integration', 'Enhanced business features on telecom-bundled plan deepens dependency'],
    ]
)

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_heading_styled('Defining Quote', level=2)
add_quote('My bank app had issues due to bad network')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run('\u2014 A simple sentence that encapsulates lost income, eroded trust, and systemic fragility. '
    'For Adesola, a network failure is not an inconvenience \u2014 it is a barrier to economic survival.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

# Save
doc_path = outdir + r'\Persona_Adekoya_Adesola.docx'
doc.save(doc_path)
print(f"Persona saved to: {doc_path}")
