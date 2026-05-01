import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)

def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)
h2 = doc.add_heading('Atolagbe Flora Yemi', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Weather-Dependent Matriarch')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201321, 2026 \u2022 6-Day Longitudinal Diary')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# ========== DIM 1 ==========
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Atolagbe Flora Yemi'],
    ['Location', 'Lagos \u2014 walks to shop daily with her children; travelled to Shagamu on one entry'],
    ['Occupation', 'Retail shop owner \u2014 physical sales, payment verification via phone'],
    ['Household', 'Married mother: husband (lives separately or travels), children, extended family network'],
    ['Diary Period', '16\u201321 April 2026 (6 entries)'],
    ['Primary Network', 'Glo (main line \u2014 identity) + MTN (data/internet \u2014 speed and affordability)'],
    ['Phone Behaviour', 'Does NOT use phone on waking (3/6 days) \u2014 domestic duties come first'],
])

doc.add_paragraph(
    'Flora is the most family-embedded respondent in this study. Her diary reads less like a '
    'work journal and more like the internal monologue of a woman who carries an entire household '
    'on her shoulders. She prays every morning (6/6 days). She cooks every day \u2014 breakfast for '
    'the family, dinner in the evening \u2014 and her evenings are spent making dinner (4/6 days). She '
    'walks to her shop with her children in tow, navigating "muddy and rough" roads while '
    'simultaneously managing a business that depends on her phone for payment verification. Her '
    'husband is mentioned only in the context of phone calls to "check up on" him, suggesting he '
    'works away from home, leaving Flora as the sole operational parent.'
)

# ========== DIM 2 ==========
add_h('Dimension 2: A Typical Day')

add_h3('Morning (06:30 \u2013 10:00)')
doc.add_paragraph(
    'Flora wakes between 6:30 and 8:00 AM, consistently starting with prayer (6/6 days). '
    'Unlike respondents who reach for their phones immediately, Flora\u2019s mornings are fully '
    'embodied: she prays, does house chores, cooks breakfast, bathes, and dresses her children. '
    'On days she uses her phone early, it is to "place a call" \u2014 specifically to "check up on '
    'husband." This is not casual scrolling; it is a relational obligation performed before the '
    'workday begins. How she ended her previous day determines how her morning starts \u2014 a '
    'pattern that reveals the cumulative toll of exhaustion.'
)
add_quote('I walked to shop with my kids. It was stressful because the road was muddy and rough due to yesterday\u2019s rain')
doc.add_paragraph(
    'The commute to her shop is on foot, with her children. On good days it is unremarkable. '
    'On rain-affected days, it becomes an ordeal. The road quality directly determines her '
    'ability to reach work and \u2014 critically \u2014 whether customers will come to her. Rain is not '
    'just weather for Flora; it is an economic event that reduces both her access to the shop '
    'and her customers\u2019 willingness to travel.'
)
add_quote('I wanted more customers but only few patronised me due to yesterday\u2019s rain')

add_h3('Afternoon (10:00 \u2013 16:00)')
doc.add_paragraph(
    'The afternoon is dominated by "patiently attending to customers" and "carefully verifying '
    'transaction alerts." Flora\u2019s primary business tool is her phone: customers pay via bank '
    'transfer, and she must verify each payment on her device. This makes her phone battery '
    'life a direct business requirement. On multiple days, she reports that "no light" prevents '
    'charging, and she explicitly names "charging my phone" as her preparation for tomorrow '
    '(3/6 days). One day, low battery prevented her from recording a journey to Shagamu \u2014 a '
    'trip she made with a friend, revealing a rare social outing.'
)
add_quote('I wanted to use my phone to record')
add_quote('I wanted to use electric cooker but I couldn\u2019t because there was no light')
doc.add_paragraph(
    'Power outages don\u2019t just affect her phone \u2014 they affect her cooking. On one day, she '
    '"couldn\u2019t cook quick because there was no light," forcing her to abandon the electric '
    'cooker. This is a uniquely gendered infrastructure burden: for Flora, power failure '
    'simultaneously threatens her business (phone dies), her domestic duty (can\u2019t cook), '
    'and her emotional state ("I will feel a little bit sad").'
)

add_h3('Evening (18:00 \u2013 22:00)')
doc.add_paragraph(
    'Evenings are for family restoration. Flora "made dinner" on 4/6 evenings \u2014 the most '
    'consistent evening activity of any respondent. Her relaxation is deeply social: she '
    '"communicated with family or friends" (3/6 days) and "watched movies" together with her '
    'family (2/6 days). She chats with friends on WhatsApp and watches comedy content. Her '
    'preparation for tomorrow is almost always the same phrase: "To charge my phone" (3/6 days). '
    'This single repeated sentence captures her entire infrastructure anxiety.'
)
add_quote('To charge my phone')

# ========== DIM 3 ==========
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Daily', 'Feed the family', 'Cooks breakfast + dinner 6/6 days; "preparing breakfast" is top priority'],
    ['Daily', 'Verify customer payments', '"carefully verify the transaction alert" \u2014 most effort-intensive task'],
    ['Relational', 'Check up on husband', 'First phone call of the day is to husband (3/6 days)'],
    ['Financial', 'Avoid unnecessary spending', '"Sales was poor and I did not want to spend the little money I made"'],
    ['Aspirational', 'Keep family safe and healthy', '"Things that keep me safe, working and healthy"'],
    ['Personal', 'Sleep and rest', '"I wanted to sleep" (3/6 days) \u2014 rest is aspirational, not guaranteed'],
])

# ========== DIM 4 ==========
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Frequency', 'Defining Quote'], [
    ['Rain \u2014 reduces patronage and mobility', '4/6 days', '"only few patronised me due to yesterday\u2019s rain"'],
    ['No light / Power outage', '4/6 days', '"no light" \u2014 prevents cooking, charging, and business'],
    ['Low battery / Charging anxiety', '3/6 days', '"To charge my phone" is her #1 prep for tomorrow'],
    ['Poor / Slow sales', '3/6 days', '"poor sales," "no sales," "slow sales"'],
    ['Poor network / Bad network', '3/6 days', '"poor network" causes stress during business hours'],
    ['Muddy roads after rain', '1/6 days', '"the road was muddy and rough" \u2014 affects walk to shop with kids'],
    ['No transportation options', '2/6 days', '"no transportation" blocks closing-time mobility'],
    ['Phone hanging / recharge failure', '1/6 days', '"I couldn\u2019t recharge from bank because my phone was hanging"'],
])
add_insight('Key Insight',
    'Flora is the most weather-sensitive respondent in the entire corpus. Rain appears as a direct '
    'stressor in 4/6 entries, affecting her in three distinct ways: (1) it makes the walk to shop '
    'with her children physically difficult, (2) it reduces customer patronage at her shop, and '
    '(3) it leaves muddy roads the following day that continue to deter foot traffic. Combined '
    'with power outages (4/6 days), Flora\u2019s business is essentially at the mercy of two '
    'environmental factors she cannot control: weather and electricity. Her J02 (Charging Anxiety) '
    'code fired 10 times \u2014 the joint highest in the study.')

# ========== DIM 5 ==========
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network (calls)', 'Glo \u2014 "that\u2019s the number I use" (identity-based loyalty)'],
    ['Secondary network (data)', 'MTN \u2014 "fast network," "affordable bundles," "the best network"'],
    ['Multi-SIM', 'YES \u2014 strategic: Glo for calls/identity, MTN for data/internet'],
    ['Phone on waking', 'Mixed \u2014 3/6 days NO (domestic duties first); 3/6 YES (call husband)'],
    ['Phone as work tool', '"verifying customers transactions and making calls," "making transactions"'],
    ['Phone for leisure', 'Minimal: comedy (2/6), TikTok (1/6), browsing (1/6). No entertainment on 3/6 days'],
    ['Phone charged on waking', 'Yes (4/6 days) but charging challenges on 4/6 days \u2014 precarious'],
])
doc.add_paragraph(
    'Flora\u2019s network loyalty is split and strategic. Her Glo line is her identity \u2014 "that\u2019s the '
    'number I use" \u2014 but she knows MTN is faster and more affordable for data. On different days '
    'she uses different networks depending on the task: Glo for personal/family calls, MTN for '
    'business internet and browsing. She even noted that calls via data are "free" and you can '
    '"chat for a long time," showing a sophisticated understanding of VoIP economics. Her phone '
    'is strictly a utility tool \u2014 entertainment usage is among the lowest in the study.'
)
add_quote('Because call is free with data and you can chat for a long time')

# ========== DIM 6 ==========
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending frequency', '4/6 days \u2014 moderate, focused on essentials'],
    ['Primary spend', 'Food for dinner \u2014 "getting food to eat," "buying food for dinner"'],
    ['Secondary spend', 'Airtime, transportation \u2014 "food and airtime" are essential'],
    ['Non-spending logic', '"Sales was poor and I did not want to spend the little money I made"'],
    ['Resourcefulness', '"We make use of the food stuffs we have at home" \u2014 avoids spending when possible'],
    ['Pain point', '"sudden increase in prices" and "low battery" blocking bank transactions'],
])
add_insight('Key Insight',
    'Flora operates on a "conservation-first" financial model. Unlike Adeola\u2019s replacement '
    'economics, Flora\u2019s strategy is to avoid spending altogether when sales are poor. She uses '
    '"food stuffs we have at home" and deliberately withholds spending on slow days. This means '
    'her financial wellbeing is entirely tethered to daily sales volume \u2014 there is no buffer, no '
    'credit facility, no savings to draw on. When she does spend, it is exclusively on food '
    'and communication (airtime/data), which she considers equally essential.')

# ========== DIM 7 ==========
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Primary method', 'In-person (family, customers at shop) + Calls (husband, friends)'],
    ['WhatsApp usage', 'Used for planning trips ("chatted with my friend about travelling to Shagamu")'],
    ['Who she contacts', 'Husband, family, friends, customers \u2014 deeply relational network'],
    ['Call duration', '5\u201342 minutes \u2014 longest calls are planned family discussions'],
    ['Conversation style', '"Discussion with family was planned while customer calls was random"'],
    ['Unique pattern', 'Gists with friends "on my way to market" \u2014 walking is social time'],
])
doc.add_paragraph(
    'Flora\u2019s communication reveals a woman whose social world is tightly woven into her daily '
    'routine. She doesn\u2019t set aside "social time" \u2014 instead, she gists with friends while walking '
    'to the market, discusses dinner plans with her husband via phone, and chats with church '
    'members in person. Her 42-minute call was a planned family discussion, showing that she '
    'invests heavily in relational maintenance. Her most revealing communication moment was '
    '"discussing breakfast plans with my family and business transactions with my customers" \u2014 '
    'family and business in the same breath, inseparable.'
)
add_quote('Discussed dinner with my husband and gisted with my friends on my way to market')

# ========== DIM 8 ==========
add_h('Dimension 8: Emotional Profile & Stress Map')

add_h3('Daily Emotional Arc')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Peaceful/Happy: Peaceful (2/6), Happy (2/6), Relaxed (1/6), Tired (1/6)', 'Starts well but tiredness from previous day carries over'],
    ['Morning', 'Purposeful and nurturing', 'Prayer, cooking, children \u2014 a deeply maternal morning ritual'],
    ['Mid-day', 'Patient but anxious', '"Patiently attending to customers" \u2014 waiting is her work'],
    ['Afternoon', 'Environmentally stressed', 'Rain, no light, low battery \u2014 external forces control her mood'],
    ['Evening', 'Warm and social', 'Dinner-making, family movies, gisting \u2014 the day\u2019s emotional reward'],
])

add_h3('Stress Triggers (Ranked)')
doc.add_paragraph('1. Rain \u2014 the single most impactful environmental factor on her livelihood')
doc.add_paragraph('2. No light / Power outage \u2014 blocks cooking, charging, and business transactions')
doc.add_paragraph('3. Poor sales \u2014 "no sales" is her most frequent emotional stressor')
doc.add_paragraph('4. Low battery anxiety \u2014 "To charge my phone" is her recurring evening mantra')
doc.add_paragraph('5. Bad network \u2014 disrupts transaction verification at the point of sale')

add_h3('Resilience Pattern')
doc.add_paragraph(
    'Flora\u2019s resilience is rooted in her family. She feels "productive and satisfied" when '
    '"meeting with family and friends" (3/6 days) and "when there\u2019s sales" (3/6 days). Her '
    'emotional baseline is "peaceful" \u2014 a word she uses twice, and which no other respondent '
    'uses at all. This peacefulness is not naivety; it is the emotional foundation of a woman '
    'who has accepted the unpredictability of her circumstances (rain, power, sales) and draws '
    'stability from the one constant she can control: her family rituals. Prayer, cooking, '
    'dinner with the children \u2014 these are her anchors when the external world fails.'
)

# ========== DIM 9 ==========
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Weather-Indexed Business Support', 'Flora loses customers every time it rains. A "rainy day bundle" \u2014 discounted data '
     'or free WhatsApp business messaging triggered during verified rainfall \u2014 would help her '
     'maintain online sales when walk-in traffic drops'],
    ['Solar Charging for Market Women', 'Her #1 preparation for tomorrow is "to charge my phone." A branded solar charging '
     'kiosk at market locations would solve her most persistent infrastructure barrier and create '
     'daily brand touchpoints'],
    ['Family Communication Bundles', 'She calls her husband daily and gists with friends on walks. A "family circle" bundle '
     'with unlimited calls to 3\u20135 designated numbers would match her deeply relational usage pattern'],
    ['SMS Transaction Alerts (Ultra-Reliable)', '"Confirming payments" is her most critical task. Transaction SMS delivery '
     'must be instant and reliable even on Glo\u2019s network. Any improvement here is directly '
     'revenue-positive for her segment'],
    ['VoIP Education Campaigns', 'She already knows that "call is free with data." Marketing that teaches traders to use '
     'WhatsApp calling to reduce airtime costs would position the brand as an ally, not just a vendor'],
])

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('I walked to shop with my kids. It was stressful because the road was muddy and rough due to yesterday\u2019s rain')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 In one sentence, Flora captures the triple burden of Lagos market motherhood: she is '
    'simultaneously a parent (walking with kids), a business owner (going to shop), and a victim '
    'of infrastructure failure (muddy roads from rain). She doesn\u2019t complain about any of these '
    'individually \u2014 she absorbs them all. The brand that lightens even one of these burdens will '
    'earn a customer who will never leave.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.save(f'{outdir}\\Persona_Atolagbe_Flora_Yemi.docx')
print("Persona saved: Atolagbe Flora Yemi (v2 deep-dive)")
