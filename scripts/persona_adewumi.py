import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)

def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)
h2 = doc.add_heading('Adewumi Omotoyosi', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Invisible Engine')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201322, 2026 \u2022 7-Day Longitudinal Diary')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# ========== DIM 1 ==========
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Adewumi Omotoyosi'],
    ['Location', 'Lagos \u2014 commutes via tricycle + BRT to Alagomeji/Yaba corridor'],
    ['Occupation', 'Remote worker (graphic design, digital tasks) + full-time homemaker'],
    ['Household', 'Multi-generational: lives with mother and young children'],
    ['Diary Period', '16\u201322 April 2026 (7 consecutive days)'],
    ['Primary Network', 'MTN (calls and data) \u2014 single SIM, no switching'],
    ['Phone Behaviour', 'Does NOT use phone immediately on waking (6/7 days)'],
])

doc.add_paragraph(
    'Omotoyosi occupies a dual role that defines her entire daily experience: she is both a '
    'remote professional (graphic designer, digital content creator) and the primary domestic '
    'operator of a multi-generational household. She cares for young children and her mother, '
    'cooking separate meals for each based on age and dietary needs. Unlike most respondents in '
    'this study, she does not reach for her phone upon waking \u2014 her first acts are always '
    'domestic: house chores, cooking, bathing. The phone enters her day only after the household '
    'is running.'
)

# ========== DIM 2 ==========
add_h('Dimension 2: A Typical Day')

add_h3('Morning (05:00 \u2013 12:00)')
doc.add_paragraph(
    'Omotoyosi wakes between 5:00 and 6:47 AM. Her morning is entirely consumed by domestic '
    'labour. Every single morning across all 7 diary days begins with house chores, cooking, '
    'and bathing \u2014 in that order. Prayer appears in 5 of 7 entries, woven into the chores '
    'rather than preceding them. Her most urgent priority is always the same: preparing '
    'breakfast. But this is not a simple task \u2014 she prepares multiple meals because her '
    'mother and children require different food.'
)
add_quote('Made breakfast for my mum and kids. These are usually different meals due to the age difference')
doc.add_paragraph(
    'On days when she leaves home (only 2 of 7 days), she takes public transport: a tricycle '
    'to the BRT terminal, then BRT to Alagomeji in Yaba. Her commute is unusually smooth \u2014 '
    '\u201ca straight 15 mins drive\u201d \u2014 making her an outlier among respondents who universally '
    'report transport stress. On the other 5 days, she stays home, operating as both caregiver '
    'and remote worker simultaneously.'
)
add_quote('Boarded a tricycle from my home to the main BRT terminal and boarded a BRT to Alagomeji bus stop in Yaba')

add_h3('Afternoon (12:00 \u2013 18:00)')
doc.add_paragraph(
    'Afternoons are where her dual life collides most visibly. On home days, she spends hours '
    'cleaning \u2014 \u201calmost 4 hrs\u201d on the children\u2019s room alone \u2014 while simultaneously submitting '
    'remote work tasks, creating graphic designs for clients and her church, and uploading '
    'social media posts. She uses her phone as her primary design tool, noting that she is '
    '\u201cnot provident with designs using the laptop.\u201d This makes her phone not just a '
    'communication device but her entire professional workstation.'
)
add_quote('I worked remotely today so I\u2019ve been able to submit my tasks. I also made some graphic designs for a client and church')
add_quote('Making the graphic designs. I\u2019m not provident with designs using the laptop')
doc.add_paragraph(
    'Her spending is minimal and entirely needs-driven. On 5 of 7 days, she spends nothing at '
    'all. When she does spend, it\u2019s only on transport (\u201cI needed to board cabs\u201d). She frames '
    'non-spending positively: \u201cI have all I need at home.\u201d'
)

add_h3('Evening (18:00 \u2013 22:00)')
doc.add_paragraph(
    'Evenings follow a rigid pattern: make dinner, then collapse. She made dinner on 6 of 7 '
    'evenings. Her relaxation is almost exclusively sleeping \u2014 not scrolling, not streaming, '
    'not gaming. On 5 of 7 days she explicitly answers \u201cnothing\u201d when asked about phone '
    'entertainment. Her relaxation time is \u201coccasionally\u201d to \u201crarely\u201d available, and she '
    'recognises this scarcity. When she does use her phone for leisure, it\u2019s movies or games '
    '\u2014 but these are rare exceptions, not habits.'
)
add_quote('Made dinner')
doc.add_paragraph(
    'She prepares clothes for tomorrow on work days (3 of 7) but not on home days, suggesting '
    'a clear mental distinction between \u201cwork mode\u201d and \u201chome mode\u201d even though both involve '
    'intense labour.'
)

# ========== DIM 3 ==========
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Daily', 'Feed the household', 'Preparing breakfast is the #1 priority in 7/7 entries'],
    ['Daily', 'Maintain household order', '4 hours cleaning children\u2019s room; supervising kids to clean'],
    ['Professional', 'Submit remote work tasks', 'Graphic design for clients + church on phone'],
    ['Relational', 'Care for mother', 'Separate meals, birthday gifts, prayer together'],
    ['Personal', 'Rest', '\u201cSleeping\u201d is her most frequent relaxation activity \u2014 rest is aspirational, not guaranteed'],
])

# ========== DIM 4 ==========
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Frequency', 'Defining Quote'], [
    ['Power outage blocking chores', '4/7 days', '\u201cPower outage. I need to iron some clothes\u201d'],
    ['Cannot do laundry', '3/7 days', '\u201cI wanted to iron and use the washing machine but couldn\u2019t. The power was interrupted\u201d'],
    ['No light', '2/7 days', '\u201cno light\u201d \u2014 prevents ironing, laundry, and electronic chores'],
    ['Bad network disrupting remote work', '1/7 days', '\u201cbad network\u201d caused stress during work'],
    ['Sleep deprivation', 'Implied', 'Wanted to sleep but couldn\u2019t (Entry 1); rarely has relaxation time'],
])
add_insight('Key Insight',
    'Omotoyosi\u2019s primary frustration is not transport or network \u2014 it is power. The J01 (Power '
    'Outage) code fired 10 times across her diary, more than any other respondent. Power outages '
    'don\u2019t just inconvenience her \u2014 they prevent her from completing domestic duties (ironing, '
    'washing machine, blender), which cascades into her already compressed schedule. She ran the '
    'generator specifically to use the blender, revealing how power gaps force expensive workarounds.')

# ========== DIM 5 ==========
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network', 'MTN \u2014 \u201cthat\u2019s the number people know me with\u201d'],
    ['Secondary network', 'None \u2014 single SIM user'],
    ['Multi-SIM', 'No'],
    ['Phone on waking', 'NO \u2014 does not use phone on waking (6/7 days)'],
    ['Phone as work tool', 'Primary design tool: graphic design, mock tests, social media management'],
    ['Phone for leisure', 'Minimal: movies (1 day), games (1 day), nothing (5 days)'],
    ['Network pain', 'Bad network causes stress during remote work (1/7 days)'],
    ['Phone charged on waking', 'No (3/7 days) \u2014 power outages affect phone charging'],
])
doc.add_paragraph(
    'Omotoyosi\u2019s phone relationship is fundamentally different from other respondents. She does '
    'not check it on waking. She does not use it for entertainment most days. Instead, her phone '
    'is a professional production tool \u2014 she creates graphic designs, submits work tasks, and '
    'manages social media content from it. Her loyalty to MTN is identity-based (\u201cthat\u2019s the '
    'number people know me with\u201d), not performance-based.'
)

# ========== DIM 6 ==========
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending frequency', '2/7 days \u2014 lowest in the entire sample'],
    ['What she spends on', 'Transport only (\u201cI needed to board cabs\u201d)'],
    ['Non-spending rationale', '\u201cI have all I need at home\u201d / \u201cI have nothing to buy\u201d'],
    ['Payment pain', 'None reported \u2014 rarely encounters digital payment friction because she rarely spends'],
    ['Provider role', 'Feeds the entire household but does not frame this as \u201cspending\u201d'],
])
add_insight('Key Insight',
    'Her near-zero spending is not frugality \u2014 it\u2019s invisibility. She provides meals, '
    'childcare, and household maintenance daily, but because these involve no cash transactions '
    '(she uses groceries already at home), her massive economic contribution is invisible in '
    'the financial data. This is a classic unpaid domestic labour pattern.')

# ========== DIM 7 ==========
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Primary method', 'WhatsApp + in-person'],
    ['Who she contacts', 'Colleagues, family (including international calls)'],
    ['Call duration', '125 minutes (international family call) \u2014 a major social event'],
    ['Conversation style', 'Mix of planned (work) and spontaneous (\u201cit just came up\u201d)'],
    ['Topics', 'Work, business, family matters'],
    ['In-person preference', '\u201cWe discussed physically\u201d \u2014 chooses face-to-face when possible'],
])
doc.add_paragraph(
    'Omotoyosi uses WhatsApp as her primary communication channel, choosing it for both work '
    'collaboration and international family calls. Her 125-minute family call was spontaneous '
    '(\u201cit just came up\u201d) and conducted via WhatsApp because \u201cit was an international call and '
    'that was the best means.\u201d This reveals WhatsApp as a cost-saving necessity, not just a '
    'convenience.'
)

# ========== DIM 8 ==========
add_h('Dimension 8: Emotional Profile & Stress Map')

add_h3('Daily Emotional Arc')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Mostly positive: Happy (3), Relaxed (3), Sick/Sleepy (1)', 'Despite heavy workload, she wakes in good spirits'],
    ['Morning', 'Calm and purposeful', 'Very few stressors reported in mornings \u2014 domestic routine provides structure'],
    ['Mid-day', 'Pressured', 'Remote work deadlines + domestic duties collide; bad network adds friction'],
    ['Afternoon', 'Exhausted', '\u201cI rested well enough this afternoon\u201d \u2014 rare and noteworthy when it happens'],
    ['Evening', 'Depleted', 'Makes dinner, then sleeps. No phone entertainment. Relaxation is \u201coccasionally\u201d to \u201crarely\u201d'],
])

add_h3('Stress Triggers (Ranked)')
doc.add_paragraph('1. Power outages blocking domestic tasks (ironing, washing, blending)')
doc.add_paragraph('2. Sleep deprivation / exhaustion')
doc.add_paragraph('3. Bad network disrupting remote work')
doc.add_paragraph('4. Children\u2019s messy rooms requiring hours of cleaning')

add_h3('Resilience Pattern')
doc.add_paragraph(
    'Omotoyosi does not complain. Her diary is remarkably free of negative language compared '
    'to other respondents. She answers \u201cNothing\u201d to stress questions far more often than anyone '
    'else \u2014 not because she has no stress, but because her stressors are so constant (cooking, '
    'cleaning, childcare, power outages) that they have become normalised. She runs the generator '
    'to use the blender. She cooks different meals for different family members. She cleans for '
    '4 hours straight. She never frames any of this as a burden \u2014 it\u2019s simply what she does.'
)

# ========== DIM 9 ==========
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Power-aware services', 'She is the most power-affected respondent (J01: 10). A telecom product that acknowledges '
     'power unreliability \u2014 e.g., ultra-low-battery mode, offline work sync \u2014 would speak directly to her reality'],
    ['Affordable data for remote work', 'She depends on mobile data for her entire professional output. Stable, affordable '
     'work-hours data plans on MTN would reduce her #1 work stressor'],
    ['Phone-based design tools', 'She creates professional graphic designs entirely on her phone. Partnerships with mobile '
     'design apps (Canva, CapCut) could position the brand as an enabler of mobile-first professionals'],
    ['International call bundles', 'Her 125-min WhatsApp international call reveals a latent need. Affordable international '
     'call bundles would earn loyalty from diaspora-connected families'],
    ['Recognition of unpaid labour', 'Marketing that acknowledges and celebrates the invisible domestic provider \u2014 the woman '
     'who feeds everyone before she feeds herself \u2014 would create deep emotional brand affinity'],
])

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('I wanted to iron and use the washing machine but couldn\u2019t. The power was interrupted early this morning')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 A sentence that reveals the daily reality of millions of Nigerian women: domestic labour '
    'is already exhausting, and infrastructure failure makes it harder. Omotoyosi doesn\u2019t ask for '
    'less work \u2014 she asks for the power to stay on.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.save(f'{outdir}\\Persona_Adewumi_Omotoyosi.docx')
print("Persona saved: Adewumi Omotoyosi")
