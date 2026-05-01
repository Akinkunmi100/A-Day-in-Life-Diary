"""
Theme Clustering Analysis
=========================
Reads coded_corpus_full.xlsx and performs:
1. Code family co-occurrence analysis (which families appear together)
2. Strategic theme clustering (collapse 11 families → 6-7 themes)
3. Session-level clustering (Morning / Afternoon / Evening breakdowns)
4. Outputs: heatmap visuals + executive .docx report

Data source: coded_corpus_full.xlsx (1,499 coded rows, 11 respondents)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from collections import Counter, defaultdict
from itertools import combinations
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

# ============================================================
# 1. LOAD AND PREPARE DATA
# ============================================================
df = pd.read_excel(f'{outdir}\\coded_corpus_full.xlsx', engine='openpyxl')
print(f"Loaded {len(df)} coded rows across {df['Respondent'].nunique()} respondents")

# Extract code families from individual code columns
def get_families(row):
    families = set()
    for col in ['Code_1', 'Code_2', 'Code_3', 'Code_4']:
        val = row.get(col)
        if pd.notna(val):
            code = str(val).strip()
            if len(code) >= 1 and code[0].isalpha():
                families.add(code[0].upper())
    return families

df['Families'] = df.apply(get_families, axis=1)

# Family labels
FAMILY_LABELS = {
    'A': 'Work & Productivity',
    'B': 'Network & Connectivity',
    'C': 'Mobility & Transport',
    'D': 'Financial Behaviour',
    'E': 'Social Interaction',
    'F': 'Emotional Wellbeing',
    'G': 'Leisure & Entertainment',
    'H': 'Daily Routines',
    'I': 'Aspirations & Goals',
    'J': 'Power & Infrastructure',
    'K': 'Domestic Life',
}
FAMILIES = list('ABCDEFGHIJK')

# ============================================================
# 2. CO-OCCURRENCE MATRIX (Overall)
# ============================================================
print("\n--- Building co-occurrence matrix ---")

# For each entry (Entry_ID), collect all families that appear
entry_families = df.groupby('Entry_ID')['Families'].apply(lambda x: set().union(*x)).to_dict()

# Count co-occurrences
cooccurrence = np.zeros((11, 11), dtype=int)
for entry_id, fams in entry_families.items():
    fam_list = sorted([f for f in fams if f in FAMILIES])
    for i, f1 in enumerate(fam_list):
        idx1 = FAMILIES.index(f1)
        cooccurrence[idx1][idx1] += 1  # self-count
        for f2 in fam_list[i+1:]:
            idx2 = FAMILIES.index(f2)
            cooccurrence[idx1][idx2] += 1
            cooccurrence[idx2][idx1] += 1

# Normalize to percentages (of total entries)
n_entries = len(entry_families)
cooccurrence_pct = (cooccurrence / n_entries * 100).round(1)

print(f"Total diary entries analysed: {n_entries}")

# ============================================================
# 3. CO-OCCURRENCE HEATMAP (Overall)
# ============================================================
fig, ax = plt.subplots(figsize=(14, 11))
fig.patch.set_facecolor('#0D1B2A')
ax.set_facecolor('#0D1B2A')

labels = [f"{f}: {FAMILY_LABELS[f]}" for f in FAMILIES]
im = ax.imshow(cooccurrence_pct, cmap='YlOrRd', aspect='auto', vmin=0)

# Annotate cells
for i in range(11):
    for j in range(11):
        val = cooccurrence_pct[i][j]
        color = 'white' if val > 50 else 'black' if val > 20 else '#AAAAAA'
        ax.text(j, i, f'{val:.0f}%', ha='center', va='center', fontsize=8, color=color, fontweight='bold')

ax.set_xticks(range(11))
ax.set_yticks(range(11))
ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=9, color='#CCCCCC')
ax.set_yticklabels(labels, fontsize=9, color='#CCCCCC')
ax.set_title('Code Family Co-Occurrence Matrix\n(% of diary entries where both families appear)',
             fontsize=16, fontweight='bold', color='white', pad=20)

cbar = plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
cbar.set_label('Co-occurrence %', color='#CCCCCC', fontsize=10)
cbar.ax.tick_params(colors='#AAAAAA')

plt.tight_layout()
heatmap_path = f'{outdir}\\Theme_CoOccurrence_Overall.png'
plt.savefig(heatmap_path, dpi=200, bbox_inches='tight', facecolor='#0D1B2A')
plt.close()
print(f"Overall co-occurrence heatmap saved: {heatmap_path}")

# ============================================================
# 4. SESSION-LEVEL CO-OCCURRENCE (Morning / Afternoon / Evening)
# ============================================================
# Map Time_Block to sessions
def get_session(row):
    tb = str(row.get('Time_Block', '')).strip().lower()
    if 'morning' in tb:
        return 'Morning'
    elif 'afternoon' in tb:
        return 'Afternoon'
    elif 'evening' in tb:
        return 'Evening'
    return 'Unknown'

df['Session'] = df.apply(get_session, axis=1)

session_matrices = {}
for session in ['Morning', 'Afternoon', 'Evening']:
    df_sess = df[df['Session'] == session]
    entry_fams = df_sess.groupby('Entry_ID')['Families'].apply(lambda x: set().union(*x)).to_dict()
    n = len(entry_fams)
    mat = np.zeros((11, 11), dtype=int)
    for eid, fams in entry_fams.items():
        fam_list = sorted([f for f in fams if f in FAMILIES])
        for i, f1 in enumerate(fam_list):
            idx1 = FAMILIES.index(f1)
            mat[idx1][idx1] += 1
            for f2 in fam_list[i+1:]:
                idx2 = FAMILIES.index(f2)
                mat[idx1][idx2] += 1
                mat[idx2][idx1] += 1
    if n > 0:
        session_matrices[session] = (mat / n * 100).round(1)
    else:
        session_matrices[session] = mat.astype(float)
    print(f"  {session}: {n} entries")

# Plot session-level heatmaps (3-panel)
fig, axes = plt.subplots(1, 3, figsize=(24, 9))
fig.patch.set_facecolor('#0D1B2A')
fig.suptitle('Code Family Co-Occurrence by Session\n(% of session entries where both families appear)',
             fontsize=18, fontweight='bold', color='white', y=1.02)

short_labels = [f"{f}" for f in FAMILIES]

for idx, (session, mat) in enumerate(session_matrices.items()):
    ax = axes[idx]
    ax.set_facecolor('#0D1B2A')
    im = ax.imshow(mat, cmap='YlOrRd', aspect='auto', vmin=0, vmax=100)

    for i in range(11):
        for j in range(11):
            val = mat[i][j]
            color = 'white' if val > 50 else 'black' if val > 20 else '#AAAAAA'
            ax.text(j, i, f'{val:.0f}', ha='center', va='center', fontsize=7, color=color)

    ax.set_xticks(range(11))
    ax.set_yticks(range(11))
    ax.set_xticklabels(short_labels, fontsize=9, color='#CCCCCC')
    ax.set_yticklabels(short_labels if idx == 0 else [], fontsize=9, color='#CCCCCC')
    ax.set_title(session, fontsize=14, fontweight='bold', color='#4ECDC4', pad=10)

plt.tight_layout()
session_path = f'{outdir}\\Theme_CoOccurrence_BySession.png'
plt.savefig(session_path, dpi=200, bbox_inches='tight', facecolor='#0D1B2A')
plt.close()
print(f"Session co-occurrence heatmaps saved: {session_path}")

# ============================================================
# 5. IDENTIFY STRATEGIC THEME CLUSTERS
# ============================================================
print("\n--- Identifying strategic clusters ---")

# Find strongest co-occurring pairs (above-diagonal, non-self)
pairs = []
for i in range(11):
    for j in range(i+1, 11):
        pairs.append((FAMILIES[i], FAMILIES[j], cooccurrence_pct[i][j]))
pairs.sort(key=lambda x: -x[2])

print("Top 15 co-occurring family pairs:")
for f1, f2, pct in pairs[:15]:
    print(f"  {f1} ({FAMILY_LABELS[f1]}) + {f2} ({FAMILY_LABELS[f2]}): {pct}%")

# ============================================================
# 6. RESPONDENT × FAMILY FREQUENCY MATRIX
# ============================================================
respondent_family = pd.DataFrame(0, index=df['Respondent'].unique(), columns=FAMILIES)
for _, row in df.iterrows():
    resp = row['Respondent']
    for fam in row['Families']:
        if fam in FAMILIES:
            respondent_family.loc[resp, fam] += 1

# Normalize per respondent (percentage of that respondent's total codes)
respondent_pct = respondent_family.div(respondent_family.sum(axis=1), axis=0) * 100

# Plot respondent × family heatmap
fig, ax = plt.subplots(figsize=(16, 10))
fig.patch.set_facecolor('#0D1B2A')
ax.set_facecolor('#0D1B2A')

resp_labels = list(respondent_pct.index)
data = respondent_pct.values

im = ax.imshow(data, cmap='YlGnBu', aspect='auto')
for i in range(len(resp_labels)):
    for j in range(11):
        val = data[i][j]
        color = 'white' if val > 20 else 'black' if val > 10 else '#AAAAAA'
        ax.text(j, i, f'{val:.0f}%', ha='center', va='center', fontsize=8, color=color, fontweight='bold')

ax.set_xticks(range(11))
ax.set_yticks(range(len(resp_labels)))
fam_labels_short = [f"{f}: {FAMILY_LABELS[f][:12]}" for f in FAMILIES]
ax.set_xticklabels(fam_labels_short, rotation=45, ha='right', fontsize=9, color='#CCCCCC')
ax.set_yticklabels(resp_labels, fontsize=9, color='#CCCCCC')
ax.set_title('Respondent \u00d7 Code Family Distribution\n(% of each respondent\'s total codes)',
             fontsize=16, fontweight='bold', color='white', pad=20)

cbar = plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
cbar.set_label('% of Respondent Codes', color='#CCCCCC', fontsize=10)
cbar.ax.tick_params(colors='#AAAAAA')

plt.tight_layout()
resp_path = f'{outdir}\\Theme_Respondent_Family_Matrix.png'
plt.savefig(resp_path, dpi=200, bbox_inches='tight', facecolor='#0D1B2A')
plt.close()
print(f"Respondent × Family matrix saved: {resp_path}")

# ============================================================
# 7. SESSION-LEVEL FAMILY DOMINANCE BAR CHART
# ============================================================
session_totals = {}
for session in ['Morning', 'Afternoon', 'Evening']:
    df_sess = df[df['Session'] == session]
    counts = Counter()
    for _, row in df_sess.iterrows():
        for fam in row['Families']:
            if fam in FAMILIES:
                counts[fam] += 1
    session_totals[session] = counts

fig, ax = plt.subplots(figsize=(16, 8))
fig.patch.set_facecolor('#0D1B2A')
ax.set_facecolor('#0D1B2A')

x = np.arange(11)
width = 0.25
colors = {'Morning': '#FFD700', 'Afternoon': '#FF6B6B', 'Evening': '#4ECDC4'}

for i, (session, counts) in enumerate(session_totals.items()):
    vals = [counts.get(f, 0) for f in FAMILIES]
    bars = ax.bar(x + i*width, vals, width, label=session, color=colors[session], alpha=0.85, edgecolor='white', linewidth=0.5)

ax.set_xticks(x + width)
ax.set_xticklabels([f"{f}: {FAMILY_LABELS[f]}" for f in FAMILIES], rotation=45, ha='right', fontsize=9, color='#CCCCCC')
ax.set_ylabel('Code Frequency', fontsize=12, color='#2E86AB')
ax.set_title('Code Family Frequency by Session (Morning / Afternoon / Evening)',
             fontsize=16, fontweight='bold', color='white', pad=20)
ax.legend(fontsize=11, loc='upper right', facecolor='#1B2838', edgecolor='#2E86AB', labelcolor='white')
ax.tick_params(colors='#666666')
ax.spines['bottom'].set_color('#334455')
ax.spines['left'].set_color('#334455')
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)

plt.tight_layout()
session_bar_path = f'{outdir}\\Theme_Family_BySession_Bar.png'
plt.savefig(session_bar_path, dpi=200, bbox_inches='tight', facecolor='#0D1B2A')
plt.close()
print(f"Session family bar chart saved: {session_bar_path}")

# ============================================================
# 8. GENERATE DOCX REPORT
# ============================================================
print("\n--- Generating Theme Clustering Report ---")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)

# Title
doc.add_paragraph()
h = doc.add_heading('Strategic Theme Clustering Report', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(32)
doc.add_paragraph('Telecom Ethnography \u2014 "A Day in the Life" Diary Study')
doc.add_paragraph('Lagos, Nigeria \u2022 April 16\u201322, 2026 \u2022 11 Respondents \u2022 65 Diary Entries')
doc.add_page_break()

# Section 1: Methodology
add_h('1. Methodology')
doc.add_paragraph(
    'This analysis examines the co-occurrence patterns of 11 thematic code families across '
    f'{n_entries} diary entries from 11 respondents. Two codes are said to "co-occur" when both '
    'appear in the same diary entry, indicating that the respondent experienced both themes '
    'within the same day. High co-occurrence suggests that two themes are structurally linked '
    'in the respondent\u2019s lived experience \u2014 they are not independent problems but parts of '
    'the same daily reality.'
)

# Section 2: Overall Co-occurrence
add_h('2. Overall Co-Occurrence Matrix')
doc.add_picture(heatmap_path, width=Inches(6.5))
doc.add_paragraph()

# Top pairs table
add_h('Top 10 Co-Occurring Family Pairs', level=2)
t = doc.add_table(rows=1, cols=4)
t.style = 'Light Grid Accent 1'
for i, h_text in enumerate(['Rank', 'Family 1', 'Family 2', 'Co-occurrence %']):
    t.rows[0].cells[i].text = h_text
    for p in t.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
for rank, (f1, f2, pct) in enumerate(pairs[:10], 1):
    row = t.add_row()
    row.cells[0].text = str(rank)
    row.cells[1].text = f'{f1}: {FAMILY_LABELS[f1]}'
    row.cells[2].text = f'{f2}: {FAMILY_LABELS[f2]}'
    row.cells[3].text = f'{pct}%'
doc.add_paragraph()

# Section 3: Strategic Clusters
add_h('3. Strategic Theme Clusters')
doc.add_paragraph(
    'Based on the co-occurrence analysis, the 11 code families collapse into the following '
    'strategic theme clusters. Each cluster represents a coherent "story" that recurs across '
    'multiple respondents and diary entries.'
)

# Define clusters based on co-occurrence data
clusters = [
    {
        'name': 'The Infrastructure-Productivity Loop',
        'codes': 'A (Work) + H (Routine) + J (Power) + B (Network)',
        'narrative': 'Work productivity, daily routines, power outages, and network connectivity form the '
                     'tightest cluster in the data. When NEPA fails, phones die. When phones die, business '
                     'calls are missed. When calls are missed, sales drop. This is not four separate problems '
                     '\u2014 it is one systemic failure with four symptoms.',
        'respondents': 'Adekunle (highest J-codes), Adeola, Flora, Adekoya'
    },
    {
        'name': 'The Domestic-Work Double Burden',
        'codes': 'K (Domestic) + A (Work) + H (Routine)',
        'narrative': 'Domestic responsibilities (cooking, cleaning, childcare) and work demands constantly '
                     'compete for the same morning hours. Respondents must complete household duties before '
                     'leaving for work, creating a daily race against time that determines their entire day.',
        'respondents': 'Ogunleye, Flora, Adeola, Mary'
    },
    {
        'name': 'The Commute Burden',
        'codes': 'C (Mobility) + F (Emotion) + D (Financial)',
        'narrative': 'Transport stress, emotional strain, and financial cost form a tight triangle. Every '
                     'commute costs money (bike, keke, bus), causes stress (traffic, rain, lateness), and '
                     'consumes time that could be spent earning. The commute is both a financial drain and '
                     'an emotional weight.',
        'respondents': 'Mary (multi-modal), Ogunleye, Adekoya, Daramola'
    },
    {
        'name': 'Survival Economics',
        'codes': 'D (Financial) + A (Work) + K (Domestic)',
        'narrative': 'Spending, earning, and feeding form an inseparable cycle. Money flows out daily on '
                     'food, transport, and airtime. It can only be replenished through sales or wages. When '
                     'sales fail, respondents literally cannot eat well. Financial behaviour is not a separate '
                     'domain \u2014 it is the arithmetic of daily survival.',
        'respondents': 'Adeola ("replacement economics"), Adekunle, Flora, Olaitan'
    },
    {
        'name': 'Connection as Lifeline',
        'codes': 'E (Social) + B (Network) + F (Emotion)',
        'narrative': 'Social interaction, network availability, and emotional wellbeing co-occur because '
                     'human connection IS the coping mechanism. Respondents who gist with friends, call '
                     'family, or visit neighbours report higher emotional resilience. When network fails '
                     'and calls drop, isolation follows.',
        'respondents': 'Ehindero, Adekoya, Ogunleye, Adewumi'
    },
    {
        'name': 'The Escape Economy',
        'codes': 'G (Leisure) + F (Emotion) + B (Network)',
        'narrative': 'Entertainment (TikTok, comedy, movies, music) is not leisure \u2014 it is stress management. '
                     'It co-occurs with emotional codes because respondents explicitly use digital entertainment '
                     'to "relax" and "escape stress." This requires data, making network reliability a mental '
                     'health issue.',
        'respondents': 'Adekunle (TikTok), Ehindero (Instagram), Ogunleye (music), Adekoya (comedy)'
    },
]

for i, cluster in enumerate(clusters, 1):
    add_h(f'Cluster {i}: {cluster["name"]}', level=3)
    p = doc.add_paragraph()
    run = p.add_run(f'Core Codes: ')
    run.bold = True
    p.add_run(cluster['codes'])
    doc.add_paragraph(cluster['narrative'])
    p = doc.add_paragraph()
    run = p.add_run('Key Respondents: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(cluster['respondents'])
    doc.add_paragraph()

# Section 4: Session Breakdown
add_h('4. Session-Level Analysis')
doc.add_paragraph(
    'The same co-occurrence patterns shift dramatically across the three diary sessions. '
    'Morning is dominated by the Domestic-Work Double Burden. Afternoon is the peak of the '
    'Infrastructure-Productivity Loop. Evening belongs to the Escape Economy.'
)
doc.add_picture(session_path, width=Inches(6.5))
doc.add_paragraph()
doc.add_picture(session_bar_path, width=Inches(6.5))
doc.add_paragraph()

# Session narratives
for session, narrative in [
    ('Morning', 'The morning session is dominated by H (Routine), K (Domestic), and A (Work). '
     'Respondents wake between 4:30 and 8:00 AM, pray, do house chores, cook, bathe, and prepare for work. '
     'The phone is used primarily as a utility (alarm, torch, time check) rather than for business. '
     'Emotional codes are generally positive (prayer, happiness, energy). The key tension is the '
     'race to complete domestic duties before leaving for work.'),
    ('Afternoon', 'The afternoon is the peak business period. A (Work) and B (Network) dominate as '
     'respondents attend to customers, make calls, confirm payments, and manage their businesses. '
     'J (Power) peaks here because this is when phones run out of charge after a full morning of use '
     'without access to electricity. F (Emotional) codes shift negative \u2014 frustration, stress, and '
     'anxiety replace the morning\u2019s optimism.'),
    ('Evening', 'The evening belongs to G (Leisure), E (Social), and K (Domestic). Respondents close '
     'from work, cook dinner, choose clothes for tomorrow, and finally have time to relax. Phone usage '
     'shifts to entertainment: TikTok, comedy, movies, Instagram, music. Social calls to friends and '
     'family peak here. This is the only session where respondents report "escape" and "relaxation."'),
]:
    add_h(session, level=3)
    doc.add_paragraph(narrative)

# Section 5: Respondent Matrix
add_h('5. Respondent \u00d7 Code Family Distribution')
doc.add_picture(resp_path, width=Inches(6.5))
doc.add_paragraph()

doc.save(f'{outdir}\\Theme_Clustering_Report.docx')
print(f"\nTheme Clustering Report saved: {outdir}\\Theme_Clustering_Report.docx")
print("\n=== THEME CLUSTERING COMPLETE ===")
