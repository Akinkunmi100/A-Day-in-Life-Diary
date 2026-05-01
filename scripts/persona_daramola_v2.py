import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)

def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('\u26a0 Methodological Note: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    p.add_run(text)

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)
h2 = doc.add_heading('Daramola Solomon', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Mobile Professional')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16, 2026 \u2022 1-Day Snapshot')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

add_note(
    'This persona is synthesised from a single diary entry (1 of a possible 7 days). While it '
    'provides a rich snapshot of Daramola\u2019s daily experience, it cannot capture day-to-day '
    'variation, emotional arcs over time, or situational differences. All frequency-based claims '
    'are drawn from this single entry only. Despite this limitation, the depth and articulacy of '
    'Daramola\u2019s responses yield valuable strategic insights \u2014 he is the most professionally-oriented '
    'respondent in the study and offers a unique perspective on white-collar Lagos life.'
)
doc.add_paragraph()

# ========== DIM 1 ==========
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Daramola Solomon'],
    ['Location', 'Lagos \u2014 short 10-minute commute via public transportation'],
    ['Occupation', 'Insurance professional / Marketer \u2014 handles policy renewals, visits brokers and underwriters'],
    ['Household', 'Appears independent; handles own chores and cooking; no family members mentioned'],
    ['Diary Period', '16 April 2026 (1 entry only)'],
    ['Primary Network', 'MTN (main line) \u2014 "that\u2019s the number people know me with"'],
    ['Phone Behaviour', 'Does NOT use phone immediately on waking \u2014 starts with bath and coffee'],
])

doc.add_paragraph(
    'Daramola is the professional outlier in this study. While most respondents are traders, '
    'shopkeepers, or informal workers, Daramola operates in the formal insurance sector \u2014 '
    'visiting "brokers and underwriters," managing "policy renewals," and conducting "marketing '
    'runs." His language is distinctly corporate: he describes his communication as "effective and '
    'less time consuming," considers "data to ensure effective service delivery" a necessity, and '
    'evaluates spending against "alternative options." He is also the only respondent who starts '
    'his day with "a cup of coffee" rather than prayer \u2014 a small but significant cultural marker '
    'that distinguishes his lifestyle from the deeply faith-anchored routines of the other 10 '
    'respondents.'
)

# ========== DIM 2 ==========
add_h('Dimension 2: A Typical Day')

add_h3('Morning (07:00 \u2013 10:00)')
doc.add_paragraph(
    'Daramola wakes at 7:00 AM \u2014 later than most respondents \u2014 and starts with a bath and his '
    'phone check. He does not use his phone immediately; instead, he reports "Didn\u2019t use it." '
    'His morning anchor is "a cup of coffee," not prayer or chores. His priority is "my work / '
    'business," and the most urgent morning task is "communication with family/customers." He '
    'commutes via public transportation for "about 10 mins drive" \u2014 the shortest reported '
    'commute in the study. During the journey, he uses his phone to "check someone I can call" '
    'and "check for information," indicating he begins work while still in transit.'
)
add_quote('It\u2019s a short journey about 10mins drive')
doc.add_paragraph(
    'Despite the short commute, rain was a stressor on this particular morning: "the rain this '
    'morning" caused delays. This is noteworthy because unlike traders whose rain-stress is '
    'about lost customers, Daramola\u2019s rain-stress is about punctuality \u2014 "keeping up with '
    'appointment time" requires precision that weather disrupts.'
)
add_quote('Yeah, the rain this morning')

add_h3('Afternoon (10:00 \u2013 16:00)')
doc.add_paragraph(
    'The afternoon is the highest-intensity period. Daramola is mobile: his work consists of '
    '"marketing runs" \u2014 visiting clients, brokers, and underwriters across Lagos to discuss '
    'policy renewals. His most effort-intensive task is "keeping up with appointment time," '
    'suggesting a schedule-driven workflow that depends on transportation reliability. New '
    'tasks interrupt his planned work: "new tasks coming in between outstanding tasks" causes '
    'stress, revealing a workload management challenge common in professional environments but '
    'rare in this respondent pool.'
)
add_quote('Attending to clients requests and following up on outstanding tasks')
add_quote('New tasks coming in between outstanding tasks')
doc.add_paragraph(
    'His phone is central to this workflow. He uses it for "calling and chatting with customers '
    'and brokers" and considers "reaching out to people" the task that would be most difficult '
    'without it. His data consumption is professional: he considers data a "necessity" for '
    '"effective service delivery," not a leisure indulgence.'
)
add_quote('Data to ensure effective service delivery')

add_h3('Evening (16:00 \u2013 22:00)')
doc.add_paragraph(
    'Daramola\u2019s evening is the most domestically self-sufficient in the study. He "made dinner" '
    'himself and plans to do "laundry and chores." He watched "videos online" and "movies" for '
    'relaxation, describing the experience as helping him "both" relax and escape stress. He '
    'is the only respondent who frames relaxation as dual-purpose ("both relax and escape '
    'stress"), suggesting a self-aware approach to emotional regulation. He did not spend time '
    'with family or friends, and reports doing "nothing" with anyone \u2014 his evening is solitary.'
)

# ========== DIM 3 ==========
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Professional', 'Meet appointment schedules', '"Keeping up with appointment time" required the most effort'],
    ['Professional', 'Maintain client relationships', '"Calling and chatting with customers and brokers"'],
    ['Professional', 'Ensure effective service delivery', '"Data to ensure effective service delivery"'],
    ['Financial', 'Optimise spending against alternatives', '"Alternative options" are considered before every purchase'],
    ['Personal', 'Relax and escape stress', 'Videos and movies serve "both" functions for him'],
])

# ========== DIM 4 ==========
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Context', 'Defining Quote'], [
    ['Task-switching overload', 'New work interrupting existing work', '"New tasks coming in between outstanding tasks"'],
    ['Cash-only payment friction', 'Some payments cannot be made digitally', '"the fact that I have to make some payments in cash"'],
    ['Rain-induced delays', 'Weather disrupting appointment punctuality', '"the rain this morning"'],
    ['Transportation stress', 'Movement across Lagos for marketing runs', '"transportation" makes him stressed/uncomfortable'],
    ['Traffic', 'Disrupts scheduled appointments', '"traffic" caused stress'],
    ['Phone charging anxiety', 'Wanted to charge but couldn\u2019t (evening)', '"Charge my phone" was a blocked goal'],
])
add_insight('Key Insight',
    'Daramola\u2019s frustration profile is fundamentally different from the trader respondents. His '
    'primary pain is not power outages or low sales \u2014 it is workflow interruption and cash-payment '
    'friction. He is the only respondent who explicitly identifies having to pay in cash as a "pain '
    'point." This reveals a professional who has fully adopted digital finance and is frustrated '
    'when forced back into the physical economy. His stress comes not from scarcity but from '
    'efficiency friction \u2014 the gap between how smoothly his day should run and how Lagos actually '
    'operates.')

# ========== DIM 5 ==========
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network', 'MTN \u2014 "that\u2019s the number people know me with"'],
    ['Secondary network', 'None mentioned \u2014 likely single-SIM'],
    ['Phone on waking', 'NO \u2014 starts with bath and coffee'],
    ['Phone as work tool', '"Calling and chatting with customers and Brokers" \u2014 the core of his job'],
    ['Phone for leisure', 'Videos online, movies \u2014 evening relaxation'],
    ['Data as necessity', '"Data to ensure effective service delivery" \u2014 professional infrastructure'],
    ['Phone charged', 'Yes on waking, but charging anxiety in evening \u2014 wanted to charge but couldn\u2019t'],
])
doc.add_paragraph(
    'Daramola\u2019s relationship with his phone is the most professionally integrated in the study. '
    'For the traders, the phone is an auxiliary tool \u2014 a calculator, a torch, a payment verifier. '
    'For Daramola, the phone IS the job. Without it, he cannot "reach out to people," cannot '
    'maintain client relationships, cannot receive incoming requests. He chose WhatsApp and calls '
    'because they are "effective and less time consuming" \u2014 an efficiency-first rationale that '
    'no other respondent articulated. His average call lasts "3 minutes," the shortest in the study, '
    'reflecting a professional brevity that maximises his appointment-packed schedule.'
)

# ========== DIM 6 ==========
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending', 'Spent money on this day \u2014 on "necessity, data, food"'],
    ['Data as investment', '"Data to ensure effective service delivery" \u2014 not a cost, an investment'],
    ['Cash friction', '"the fact that I have to make some payments in cash" \u2014 prefers digital'],
    ['Decision-making', '"Alternative options" are evaluated before spending'],
    ['Professional spending', 'Data is prioritised alongside food as essential'],
])
add_insight('Key Insight',
    'Daramola is the only respondent who frames data as a professional investment rather than a '
    'personal expense. While traders buy airtime to "call customers," Daramola buys data to '
    '"ensure effective service delivery." This language reveals a mindset where connectivity is '
    'not consumption \u2014 it is revenue infrastructure. His frustration with cash payments suggests '
    'he is an early adopter of digital finance who is held back by a cash-dependent ecosystem.')

# ========== DIM 7 ==========
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Primary method', 'Calls + WhatsApp + In-person \u2014 "effective and less time consuming"'],
    ['Who he contacts', 'Colleagues and Customers \u2014 purely professional on this day'],
    ['Call duration', '"An average of 3 minutes" \u2014 the shortest in the study'],
    ['Topics', '"Work, policy renewals" \u2014 highly focused'],
    ['Volume', 'Multiple short calls throughout the day \u2014 high-frequency, low-duration'],
    ['WhatsApp rationale', '"makes communication easy" \u2014 efficiency-driven'],
])
doc.add_paragraph(
    'Daramola\u2019s communication style is "professional burst" \u2014 high-frequency, short-duration calls '
    'and messages throughout the day. His 3-minute average call length contrasts sharply with Flora\u2019s '
    '42-minute family discussions or Adeola\u2019s 2-hour "life and business" chats. Every communication '
    'is purposeful and time-bounded. He views WhatsApp not as a social platform but as a professional '
    'CRM tool \u2014 a channel for coordinating appointments, confirming renewals, and managing his '
    'client pipeline.'
)

# ========== DIM 8 ==========
add_h('Dimension 8: Emotional Profile & Stress Map')

add_h3('Single-Day Emotional Arc')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Tired', '"Tired" \u2014 the only respondent who starts the day with explicit fatigue'],
    ['Morning', 'Pragmatic', 'Bathes, coffee, commutes \u2014 functional without emotional language'],
    ['Mid-day', 'Pressured but productive', '"Attending to clients requests" \u2014 purposeful stress'],
    ['Afternoon', 'Stressed by interruptions', '"New tasks coming in between outstanding tasks"'],
    ['Evening', 'Self-regulated', 'Videos for "both" relaxation and stress escape \u2014 deliberate decompression'],
])

add_h3('Stress Triggers')
doc.add_paragraph('1. Workflow interruption \u2014 new tasks disrupting planned schedule')
doc.add_paragraph('2. Cash-payment requirements \u2014 forced back into physical finance')
doc.add_paragraph('3. Rain and transportation \u2014 threatening appointment punctuality')
doc.add_paragraph('4. Phone battery \u2014 charging anxiety in the evening')

add_h3('Resilience Pattern')
doc.add_paragraph(
    'Daramola is the most self-aware emotional regulator in the study. He explicitly states that '
    'watching videos helps him "both relax and escape stress" \u2014 a dual-function understanding '
    'that suggests deliberate self-care. He wakes tired but pushes through with coffee (not '
    'prayer, not chores \u2014 caffeine as fuel). His satisfaction trigger is identical to the traders: '
    '"when there\u2019s sales." Despite his professional veneer, his emotional state is ultimately '
    'indexed to revenue \u2014 the same fundamental driver as every other respondent in the study.'
)

# ========== DIM 9 ==========
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Professional Data Plans', 'Daramola views data as "effective service delivery" infrastructure. A business-tier '
     'data plan on MTN \u2014 with guaranteed speed during work hours, priority network access, and '
     'WhatsApp Business integration \u2014 would match his professional identity'],
    ['Digital Payment Evangelism', 'He is frustrated by cash payments. Partnering with fintechs to reduce cash '
     'dependency for his insurance transactions would solve his #1 financial pain point and '
     'position the brand as a catalyst for professional digitisation'],
    ['Appointment Management Tools', '"Keeping up with appointment time" is his biggest challenge. A simple SMS-based '
     'appointment reminder service (broadcast to clients and self) would reduce his scheduling stress'],
    ['Evening Streaming Bundles', 'He watches videos nightly for emotional regulation. A post-7pm streaming bundle with '
     'high-speed data at reduced cost would serve his decompression pattern'],
    ['White-Collar Community Building', 'Daramola is isolated \u2014 no family or friend interaction on this day. A professional '
     'networking feature or business community forum via the telecom platform could address the '
     'social gap in his life'],
])

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('The fact that I have to make some payments in cash')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 In a study dominated by traders who struggle to access digital finance, Daramola is the '
    'reverse: a professional who has embraced digital payments and is frustrated when the world '
    'forces him backward into cash. His pain point is not "I can\u2019t access the bank app" \u2014 it is '
    '"I shouldn\u2019t have to use cash at all." He represents the future of Lagos commerce, trapped '
    'in the infrastructure of the present. The brand that helps him complete the digital transition '
    'will earn the loyalty of every upwardly mobile professional in the city.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.save(f'{outdir}\\Persona_Daramola_Solomon.docx')
print("Persona saved: Daramola Solomon (v2 deep-dive)")
