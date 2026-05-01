import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs: run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs: run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs: run.bold = True; run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(10)
    doc.add_paragraph()
def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True; run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

# TITLE
doc.add_paragraph(); doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs: run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F); run.font.size = Pt(36)
h2 = doc.add_heading('Olaitan Toyeeb', level=1)
for run in h2.runs: run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB); run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Exhausted Educator')
run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66); run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201322, 2026 \u2022 7-Day Longitudinal Diary')
run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# DIM 1
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Olaitan Toyeeb'],
    ['Location', 'Lagos \u2014 home-based teacher, travels to market and teaching sites'],
    ['Occupation', 'Teacher (Islamic ethics, academic subjects for children) + small-scale trader'],
    ['Household', 'Lives with wife and young daughter'],
    ['Diary Period', '16\u201322 April 2026 (7 consecutive days)'],
    ['Primary Network', 'MTN (single SIM) \u2014 \u201cIt is the best network\u201d'],
    ['Transport', 'Owns a bike/motorbike \u2014 spends on fuel, not public transport'],
])
doc.add_paragraph(
    'Toyeeb is a young family man who splits his day between teaching children and running '
    'a small trading business. He teaches Islamic ethics and academic subjects from home, '
    'using his phone to access teaching materials. Unlike most respondents who rely on '
    'public transportation, he owns a bike and spends on fuel \u2014 giving him more mobility '
    'control but creating a different financial pressure. He is the most consistently '
    'fatigued respondent in the study, reporting \u201cTired\u201d as his morning feeling in 3 of 7 days.'
)

# DIM 2
add_h('Dimension 2: A Typical Day')
add_h3('Morning (05:00 \u2013 12:00)')
doc.add_paragraph(
    'Toyeeb wakes between 6:00 and 8:30 AM \u2014 the latest waker in the study. His morning is '
    'determined by \u201chow I ended my previous day,\u201d and he frequently reports being too tired to '
    'wake on time. He begins with prayer and house chores, but prayer is his true anchor: on '
    'days when he oversleeps, \u201csaying my morning prayers\u201d is still the first thing he handles. '
    'He does NOT use his phone on waking (7/7 days), and his phone is frequently uncharged '
    '(4/7 days) due to power challenges.'
)
add_quote('I wanted to wake up earlier, I couldn\u2019t because I felt so tired')
add_quote('In the last two hours, I was still asleep. I woke up more than 30 minutes ago. Did my morning prayer and got ready to leave')
doc.add_paragraph(
    'His mornings are consumed by teaching: \u201cI started my class with my students as early as 9am. '
    'I\u2019ve been in the class with them since then.\u201d He teaches from home on most days, using his '
    'phone to access Islamic and academic teaching materials. On days he leaves home, he rides his '
    'bike to teaching sites or the market.'
)

add_h3('Afternoon (12:00 \u2013 18:00)')
doc.add_paragraph(
    'Afternoons are a continuation of the teaching grind. His most repeated answer to \u201cwhat took '
    'most of your time today?\u201d is simply: \u201cTeaching.\u201d He uses his phone as a critical teaching '
    'tool \u2014 accessing lesson materials, checking on student progress, and communicating with '
    'customers for his trading side-hustle. His conversations with family, friends, and colleagues '
    'are brief (30 minutes on average) and always spontaneous (\u201cit came up\u201d), not planned.'
)
add_quote('I wanted to have free time, but I couldn\u2019t have it because I had to go out to teach')
doc.add_paragraph(
    'His financial life reveals a man stretched thin. He spends on transport fuel and family needs, '
    'and on one revealing day, he \u201cwanted to eat well\u201d but \u201ccouldn\u2019t because I needed to spend '
    'little.\u201d He has considered borrowing: \u201cShould I keep the money for another time or take a '
    'loan from my friends to clear what\u2019s on me.\u201d This is the clearest signal of financial '
    'precarity in the entire study.'
)
add_quote('I wanted to eat well. I couldn\u2019t because I needed to spend little')
add_quote('Should I keep the money for another time or take a loan from my friends to clear what\u2019s on me')

add_h3('Evening (18:00 \u2013 22:00)')
doc.add_paragraph(
    'Evenings are sparse. He makes dinner on some days and closes from work on others. His '
    'preparation for tomorrow is always the same: \u201cTo charge my phone\u201d (3/7 entries) or \u201cto sleep.\u201d '
    'Relaxation is rare \u2014 he reports \u201cRarely\u201d or \u201cOccasionally\u201d for how often he gets downtime. '
    'When he does relax, it\u2019s always sports content on MTN. Sports is his only leisure interest '
    'across all 7 days, making him the most single-content respondent in the study.'
)

# DIM 3
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Daily', 'Teach his students effectively', '\u201cTeaching\u201d is the dominant activity in 7/7 entries'],
    ['Daily', 'Support his family before leaving home', '\u201cHelping my family at home before leaving for work\u201d'],
    ['Financial', 'Stretch limited funds', 'Considered borrowing from friends; restricted eating to save money'],
    ['Personal', 'Get enough sleep', '\u201cI wanted to wake up earlier\u201d but couldn\u2019t due to exhaustion \u2014 3/7 days'],
    ['Relational', 'Spend time with wife and daughter', '\u201cI spent some time with my wife and daughter. I watched some TikTok videos\u201d'],
])

# DIM 4
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Frequency', 'Defining Quote'], [
    ['Chronic fatigue / exhaustion', '5/7 days', '\u201cFatigue from the weekend duties\u201d'],
    ['No free time', '4/7 days', '\u201cI wanted to have free time, but I couldn\u2019t\u201d'],
    ['Power / charging challenges', '4/7 days', 'Phone uncharged on 4 mornings; \u201clight issues during this hour\u201d'],
    ['Financial strain', '2/7 days', '\u201cI needed to spend little\u201d; considering loans from friends'],
    ['Cannot eat properly', '2/7 days', '\u201cI wanted to eat this morning before attending to students, but I couldn\u2019t\u201d'],
    ['Market trips as stressor', '2/7 days', '\u201cgoing to the market\u201d caused stress'],
])
add_insight('Key Insight',
    'Toyeeb\u2019s pain is not about broken infrastructure or bad networks \u2014 it is about exhaustion. '
    'He is tired every day, cannot eat when he wants to, cannot wake when he needs to, and '
    'has no free time. His phone\u2019s battery reflects his own state: frequently depleted, struggling '
    'to recharge. He is the respondent most likely to churn from any service that adds cognitive '
    'load rather than reducing it.')

# DIM 5
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network', 'MTN \u2014 \u201cIt is the best network\u201d / \u201cthat\u2019s the number I use\u201d'],
    ['Multi-SIM', 'No \u2014 single SIM, single network'],
    ['Phone on waking', 'NO (7/7 days) \u2014 never uses phone first thing'],
    ['Phone charged on waking', 'No (4/7 days) \u2014 chronic power issues'],
    ['Phone as work tool', 'Teaching material access; customer demand tracking'],
    ['Phone as alarm', '\u201cI would wake late\u201d without phone \u2014 phone serves as alarm clock'],
    ['Phone for leisure', 'Sports content only \u2014 browsing internet, sports highlights'],
    ['Preparation for tomorrow', '\u201cTo charge my phone\u201d \u2014 charging is an aspiration, not a guarantee'],
])
doc.add_paragraph(
    'Toyeeb\u2019s relationship with his phone is functional and utilitarian. He does not scroll on '
    'waking. He does not use it for social media. His phone serves three purposes: teaching tool, '
    'alarm clock, and sports content viewer. His MTN loyalty is simple and unconditional \u2014 he '
    'doesn\u2019t compare networks or consider switching. The most telling detail: when asked how he '
    'prepares for tomorrow, his answer is not \u201cchoose clothes\u201d or \u201cplan tasks\u201d \u2014 it\u2019s \u201cto '
    'charge my phone.\u201d Charging is not a given; it is a task that requires planning.'
)

# DIM 6
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending frequency', '4/7 days'],
    ['What he spends on', 'Transport fuel (bike), food, family needs'],
    ['Budget pressure', 'Considers borrowing from friends; restricts food to save money'],
    ['Payment pain', 'None reported \u2014 transactions are cash-based'],
    ['Decision logic', '\u201cIf there\u2019s a possible way not to spend at that moment\u201d \u2014 avoidance-first mentality'],
    ['Provider role', 'Spends on family before himself; helping family \u201cwas necessary\u201d'],
])

# DIM 7
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Methods', 'Calls + WhatsApp + In-person (all three on 3/7 days)'],
    ['Who', 'Family, friends, colleagues, customers'],
    ['Topics', 'Business, football, academy, food, teaching'],
    ['Call duration', 'Short: \u201cfew minutes\u201d to \u201c30 minutes\u201d to \u201c1 hour\u201d'],
    ['Planned vs spontaneous', 'Always spontaneous: \u201cit came up\u201d (4/4 entries)'],
    ['Satisfaction', '\u201cHelping my colleagues or family\u201d makes him feel productive'],
])

# DIM 8
add_h('Dimension 8: Emotional Profile & Stress Map')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Tired (3 days), Energetic (1), missing (3)', 'Chronic fatigue dominates mornings'],
    ['Morning', 'Purposeful but pressured', 'Teaching starts early; can\u2019t complete personal tasks'],
    ['Afternoon', 'Stressed by work volume', '\u201cwork\u201d is the stressor in 4/7 entries'],
    ['Evening', 'Too tired to relax', '\u201cRarely\u201d gets relaxation time; \u201cworking\u201d even in evening'],
])
doc.add_paragraph(
    'Toyeeb\u2019s emotional profile is defined by a single word: depletion. He starts tired, works '
    'through the fatigue, and ends too exhausted to enjoy his evening. His satisfaction comes from '
    'helping others \u2014 colleagues, family, students \u2014 not from personal achievement. His stress '
    'comes from work volume, not from infrastructure failure. He is the most internally-driven '
    'stressor in the study: his burden is self-imposed obligation, not external friction.'
)

# DIM 9
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Ultra-low-cost teaching data plans', 'He depends on mobile data for teaching materials. An education-specific '
     'data plan on MTN would directly support his livelihood'],
    ['Battery/power solutions', 'His phone is uncharged 4/7 mornings. Power bank bundles or charging station '
     'partnerships would solve his #1 preparation concern'],
    ['Sports content bundles', 'Sports is his ONLY leisure content. Affordable sports streaming partnerships '
     'would create deep engagement with his single relaxation habit'],
    ['Simple, low-friction services', 'He is too tired for complex interfaces. Any product that adds steps or '
     'cognitive load will be abandoned. Simplicity = loyalty'],
    ['Family micro-loans or savings', 'He considers borrowing from friends. A telecom-linked micro-savings or '
     'micro-loan product would meet an unmet financial need'],
])

# DEFINING QUOTE
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('I wanted to eat well. I couldn\u2019t because I needed to spend little')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 A man who feeds others before himself, teaches others before resting, and charges his '
    'phone before he charges his own energy. Toyeeb\u2019s diary is a portrait of quiet sacrifice.')
run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66); run.italic = True

doc.save(f'{outdir}\\Persona_Olaitan_Toyeeb.docx')
print("Persona saved: Olaitan Toyeeb")
