import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

def add_table_styled(rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

doc = Document()

# Adeola Olowolagba Persona
add_heading_styled('Customer Persona: Adeola Olowolagba', 0)
doc.add_paragraph('Persona Archetype: The Financial Juggler / "The Shopkeeper Mother"')
doc.add_paragraph('Cluster Codes: A01, A05, B06, C01, D02, F03, H01, I01, J02, K03')

doc.add_paragraph('---')

# Dimension 1
add_heading_styled('Dimension 1: Demographics & Life Context')
add_table_styled([
    ['Gender', 'Female'],
    ['Occupation', 'Shop Owner / Trader (Attends to physical and online customers)'],
    ['Location Context', 'Lagos (Commutes to shop; manages household in Lagos)'],
    ['Financial Context', 'Extremely budget-conscious; weighs every spend against school fees and rent.'],
    ['Key Goal', 'Increasing sales and business growth while ensuring her children’s education and household needs are met.'],
    ['Household Role', 'Working Mother; responsible for children, cooking, and house chores.']
], ['Field', 'Detail'])

# Dimension 2
add_heading_styled('Dimension 2: A Typical Day Summary')
doc.add_paragraph(
    'Adeola’s day is a perpetual cycle of household labor and business hustle. '
    'She wakes between 6:00 AM and 8:00 AM, anchoring her morning in prayer (H01) '
    'before diving into chores and preparing her children for school. Her workday '
    'is spent "at shop," waiting for and attending to customers. She is heavily '
    'dependent on her phone for "online customers" and "confirming payments," '
    'making power outages a significant professional threat. Her day ends with '
    'further chores and decompression through digital entertainment.'
)

doc.add_paragraph('**The Morning (6:00 AM - 10:00 AM)**')
doc.add_paragraph(
    'Mornings are "mood-dependent" but always labor-intensive. After prayer, '
    'she "tidies up the house" and "takes children to school." She checks her '
    'phone to confirm the time and look for customer messages. Traffic is a '
    'recurring obstacle, often delaying her arrival at the shop beyond her 10:00 AM goal.'
)

doc.add_paragraph('**The Afternoon (12:00 PM - 4:00 PM)**')
doc.add_paragraph(
    'The core work period is spent "waiting for customers" and "shading market." '
    'She is a strategic business owner, using her phone to receive calls and '
    'manage online orders. A major pain point is "no light" at the shop, '
    'leading to "low battery" and the fear that customers cannot reach her.'
)

doc.add_paragraph('**The Evening (6:00 PM - 10:00 PM)**')
doc.add_paragraph(
    'Closing from work leads to "more chores" (laundry/cooking). She finds time '
    'to relax by "watching films on my phone" (G01) or "comedy" on WhatsApp/YouTube. '
    'Preparation for the next day includes choosing work clothes.'
)

# Dimension 3
add_heading_styled('Dimension 3: Frustrations & Pain Points')
add_table_styled([
    ['Category', 'Specific Pain Point', 'Impact'],
    ['Financial', 'Lack of money / Low sales (D02)', 'Forces her to "spend little" and even skip meals.'],
    ['Infrastructure', 'No light / Charging challenges (J01)', 'Disrupts business communication with customers.'],
    ['Physical', 'Chronic tiredness from chores (F08)', 'Leads to "unmet needs" for rest and sleep.'],
    ['Mobility', 'Traffic delays (C01)', 'Prevents her from reaching her shop on time.']
], ['Category', 'Detail', 'Consequence'])

# Dimension 4
add_heading_styled('Dimension 4: Network & Tech Behaviour')
doc.add_paragraph(
    'Adeola is a Multi-SIM user (B06). Airtel is her "main line" for calling because '
    '"that’s the number people know," while MTN is her preferred network for '
    '"fast" data and internet access. She relies on her phone for "confirming payments" '
    'and accessing "online customers."'
)

# Dimension 5
add_heading_styled('Dimension 5: Finances')
doc.add_paragraph(
    'She is a master of financial trade-offs. She prioritizes "school fees" and "house rent" '
    'over her own immediate comfort (I01). She views daily food spend as a "must," '
    'but weighs every other expense against the need to "replace the money spent" via sales.'
)

# Dimension 6
add_heading_styled('Dimension 6: Communication Profile')
doc.add_paragraph(
    'Business-centric communication. She interacts primarily with "customers" to discuss '
    '"market prices and availability." She uses calls and WhatsApp for "easy communication."'
)

# Dimension 7
add_heading_styled('Dimension 7: Emotional Profile')
doc.add_paragraph(
    'Adeola’s mood is linked to her financial state: she feels "productive and satisfied" '
    'specifically "when I get credited." She experiences frustration when goals are '
    'blocked by "no money" or "low sales." Her resilience is maintained through '
    'her spiritual practice and evening decompression.'
)

# Dimension 8
add_heading_styled('Dimension 8: Thematic Code Profile')
doc.add_paragraph('**Dominant Themes:**')
doc.add_paragraph('1. **The Hustling Shopkeeper**: Business as a constant, multi-channel effort.')
doc.add_paragraph('2. **The Financial Juggler**: Survival through extreme budget discipline.')
doc.add_paragraph('3. **Infrastructure Anxiety**: The professional threat of no power.')

# Dimension 9
add_heading_styled('Dimension 9: Brand Opportunities')
doc.add_paragraph(
    '1. **Trade-Credit / Small Business Loans**: Helping shopkeepers bridge "low sales" periods.\n'
    '2. **Portable Charging for Traders**: Affordable, rugged power solutions for market shops.\n'
    '3. **Budget Management Apps**: Tools that help manage the trade-off between school fees/rent and daily spend.'
)

# Save
doc_path = f'{outdir}\\Persona_Adeola_Olowolagba.docx'
doc.save(doc_path)
print(f"Saved: {doc_path}")
