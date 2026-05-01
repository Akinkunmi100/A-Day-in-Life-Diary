import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

outdir = r'c:\Users\LUMEN GLOB AL\Documents\TAOF\Thematic Analysis'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

def add_table(hdrs, rows_data):
    t = doc.add_table(rows=1+len(rows_data), cols=len(hdrs))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'\u201c{text}\u201d')
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)

def add_insight(label, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    p.add_run(text)

# ========== TITLE PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
h = doc.add_heading('Customer Persona', level=0)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x6F)
    run.font.size = Pt(36)
h2 = doc.add_heading('Adekunle Adepetun', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(28)
p = doc.add_paragraph()
run = p.add_run('Persona Archetype: The Power-Obsessed Shopkeeper')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Telecom Ethnography Project \u2014 \u201cA Day in the Life\u201d Diary Study')
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run('Lagos, Nigeria \u2022 April 16\u201321, 2026 \u2022 6-Day Longitudinal Diary')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_page_break()

# ========== DIM 1 ==========
add_h('Dimension 1: Demographics & Life Context')
add_table(['Field', 'Detail'], [
    ['Name', 'Adekunle Adepetun'],
    ['Location', 'Lagos \u2014 shop is at his residence ("my shop is at where I stay"); walks everywhere'],
    ['Occupation', 'Shop owner / Retail trader \u2014 sells goods, serves walk-in and phone customers'],
    ['Household', 'Likely head of household; interacts with neighbours, church members, CDA chairman'],
    ['Diary Period', '16\u201321 April 2026 (6 entries)'],
    ['Primary Network', 'MTN (main line) \u2014 single-network loyalty, "fast network"'],
    ['Phone Behaviour', 'Uses phone immediately on waking (5/6 days) \u2014 primarily to turn off alarm'],
])

doc.add_paragraph(
    'Adekunle is the earliest riser in the study \u2014 waking at 4:30 AM on 3 of 6 days. His shop is '
    'at his home, which eliminates commute stress but creates a different problem: his work and '
    'personal life are physically inseparable. He prays every morning (6/6 days), does house chores '
    '(5/6 days), and then transitions directly into "the shop." But his diary reveals a man whose '
    'daily life is dominated by two forces he cannot control: power outages and government council '
    'officials. His phone is not a luxury \u2014 it is a calculator, a torch, a crisis hotline, and '
    'his only connection to customers when he cannot see them face-to-face. When his phone dies, '
    'his entire world contracts to the four walls of his shop.'
)

# ========== DIM 2 ==========
add_h('Dimension 2: A Typical Day')

add_h3('Morning (04:30 \u2013 09:00)')
doc.add_paragraph(
    'Adekunle\u2019s mornings begin in the dark. He wakes at 4:30\u20136:00 AM, and his first act is turning '
    'off his phone alarm (3/6 days). He then prays \u2014 the "Holy Spirit / Prayer" determines his '
    'morning on 4/6 days. House chores follow: prayer, chores, bath, then straight to the shop. '
    'Because his shop is at his home, there is no commute \u2014 "I don\u2019t need to trek because my shop '
    'is at where I stay." On church days (2/6), he walks to church, checking the time on his phone '
    'along the way so he can "rush" if running late.'
)
add_quote('I don\u2019t need to trek because my shop is at where I stay')
doc.add_paragraph(
    'His morning priority is always sales-focused: "To make sales" is his stated concern on 3/6 '
    'days. He feels "so happy" (3/6) or "relaxed" (2/6) on waking, but his emotional state darkens '
    'rapidly when he encounters the day\u2019s first obstacle \u2014 usually a dead phone or a power outage.'
)

add_h3('Afternoon (09:00 \u2013 16:00)')
doc.add_paragraph(
    'The afternoon is where Adekunle\u2019s two persistent stressors converge. On the business side, he '
    'spends hours "sitting in the shop," "awaiting customers," and "selling market." His phone is '
    'his calculator ("I used it to calculate something"), his torch ("put-on tourch"), and his '
    'payment communicator ("to communicate with my customers"). On one critical day, government '
    'council workers confronted him over a "shop permit issue," consuming his entire day\u2019s energy '
    'and forcing him to spend money settling the dispute. He used his phone to "call chairman to '
    'mediate on the issue" \u2014 revealing the phone as a crisis management tool, not just a business aid.'
)
add_quote('settled shop permit issue with government workers (council)')
add_quote('I used it to call chairman to mediate on the issue')
doc.add_paragraph(
    'But the defining afternoon experience is power. On one devastating day (Entry 5), his entire '
    'diary is a chronicle of a phone dying: "In the shop with flat battery," "To charge my phone" '
    '(wanted but couldn\u2019t, 4 consecutive time blocks), "Nepa, so I can charge," "finding how to '
    'charge my phone." He eventually spent money on a bike ride "looking for where to charge my '
    'phone." His power bank was "faulty," NEPA (the power utility) provided no electricity, and '
    'his phone was completely dead. On this day, he could only communicate with his neighbour '
    '"in person" because his "phone was down."'
)
add_quote('In the shop with flat battery')
add_quote('I used it for bike looking for where to charge my phone')

add_h3('Evening (16:00 \u2013 22:00)')
doc.add_paragraph(
    'Evenings are for closing shop and minimal recovery. He sleeps (4/6 days) and browses TikTok '
    '(3/6 days) as his primary relaxation. On days when his phone is alive, he uses the internet '
    'to escape stress \u2014 TikTok is his consistent choice. On days when his phone is dead, he does '
    '"nothing" for entertainment and goes to sleep. His preparation for tomorrow is either choosing '
    'clothes (2/6) or \u2014 revealingly \u2014 "To charge my phone" (2/6) or simply "to sleep" (2/6).'
)
add_quote('I wanted to eat well. I couldn\u2019t because I needed to spend little')

# ========== DIM 3 ==========
add_h('Dimension 3: Goals & Motivations')
add_table(['Type', 'Goal', 'Evidence'], [
    ['Daily', 'Make sales', '"To make sales" \u2014 stated priority on 3/6 days'],
    ['Daily', 'Keep phone charged', '"To charge my phone" / "To see my phone charged" \u2014 a daily battle'],
    ['Business', 'Protect the shop from regulatory threats', 'Spent an entire day on "shop permit issue with government council"'],
    ['Relational', 'Maintain community ties', 'Interacts with neighbours, church members, CDA chairman, colleagues'],
    ['Personal', 'Escape stress via TikTok', 'TikTok is his consistent digital retreat (3/6 days)'],
    ['Personal', 'Eat well', '"I wanted to eat well" but couldn\u2019t afford it \u2014 an unmet aspiration'],
])

# ========== DIM 4 ==========
add_h('Dimension 4: Frustrations & Pain Points')
add_table(['Pain Point', 'Frequency', 'Defining Quote'], [
    ['Power outage / No light / NEPA', '5/6 days', '"poor power supply," "no light," "Nepa, so I can charge"'],
    ['Dead phone / Flat battery', '4/6 days', '"In the shop with flat battery"'],
    ['Faulty power bank', '2/6 days', '"My power bank was faulty," "My power bank because our light has issue"'],
    ['Low sales / No sales', '4/6 days', '"Low sales" \u2014 recurring morning-through-afternoon stressor'],
    ['Government council dispute', '1/6 days (3 time blocks)', '"Issue with government council" \u2014 slowed him down 3 consecutive blocks'],
    ['Rain affecting shop', '2/6 days', '"The rain affected my shop"'],
    ['Hunger during work', '2/6 days', '"I am feeling hungry," "I wanted to eat well but couldn\u2019t"'],
])
add_insight('Key Insight',
    'Adekunle has the highest J (Power & Infrastructure) code count in the entire study: 30 codes, '
    'with J02 (Charging Anxiety) firing 18 times. His diary is essentially a power-crisis journal. '
    'On Entry 5, the word "charge" appears in his response to 6 different questions. He spent '
    'money on a motorbike taxi specifically to find somewhere to charge his phone. His power bank '
    'was broken. NEPA provided no electricity. He could not use his phone for business AT ALL that '
    'day. This is not an inconvenience \u2014 this is a man whose entire livelihood was shut down by '
    'a dead battery. The cascading effect: no power \u2192 no phone \u2192 no calculator \u2192 no customer '
    'communication \u2192 no sales \u2192 can\u2019t eat well \u2192 goes to sleep hungry.')

# ========== DIM 5 ==========
add_h('Dimension 5: Phone & Network Relationship')
add_table(['Aspect', 'Detail'], [
    ['Primary network', 'MTN \u2014 "that\u2019s the number people know me with" and "fast network"'],
    ['Secondary network', 'None \u2014 single-SIM user, fully MTN-loyal'],
    ['Multi-SIM', 'No'],
    ['Phone on waking', 'YES \u2014 turns off alarm (3/6), checks time (1/6), calls customer (1/6)'],
    ['Phone as work tool', 'Calculator, torch, customer communicator, crisis hotline (CDA chairman)'],
    ['Phone for leisure', 'TikTok (3/6 days), browsing internet (3/6 days), nothing when battery dead'],
    ['Phone charged on waking', 'No (2/6 days) \u2014 power challenges on 4/6 days'],
    ['Without phone', '"I will feel a little bit sad" (3/6 days), "I would wake late" (1/6)'],
])
doc.add_paragraph(
    'Adekunle\u2019s phone is a Swiss Army knife for survival. He uses it as a calculator in his shop '
    '("for calculation"), a flashlight when there\u2019s no power ("put-on tourch"), a clock and alarm, '
    'and a crisis communication tool (calling the CDA chairman during the council dispute). When '
    'asked what he\u2019d do without his phone, he says he\u2019d "feel a little bit sad" \u2014 an understatement '
    'given that one day without a charged phone meant he couldn\u2019t communicate, couldn\u2019t calculate, '
    'and had to communicate only "in person" with his immediate neighbour. His phone IS his business '
    'infrastructure.'
)

# ========== DIM 6 ==========
add_h('Dimension 6: Financial Behaviour')
add_table(['Aspect', 'Detail'], [
    ['Spending frequency', '6/6 days \u2014 spends every single day'],
    ['Primary spend', 'Basic survival: "pure water," "drink," "water to drink"'],
    ['Notable spend', '"Settle shop permit issue" \u2014 forced government/regulatory spend'],
    ['Church spend', '"Offering" \u2014 religious obligation spending'],
    ['Charging spend', '"bike looking for where to charge my phone" \u2014 infrastructure spend'],
    ['Spending logic', '"My business" is always the consideration before spending'],
    ['Food anxiety', '"I wanted to eat well. I couldn\u2019t because I needed to spend little"'],
])
add_insight('Key Insight',
    'Adekunle spends money every single day (6/6) \u2014 but his spending is almost entirely on bare '
    'survival: water, drinks, a church offering, and a bike ride to find electricity. The most '
    'revealing spend is the bike taxi to charge his phone. When infrastructure fails, he is forced '
    'to convert the problem into a cash expense \u2014 paying for transportation to solve a power '
    'problem. This "infrastructure tax" is an invisible cost that traditional financial analysis '
    'would miss. He also cannot afford to eat well, suggesting his daily revenues barely cover '
    'survival needs.')

# ========== DIM 7 ==========
add_h('Dimension 7: Communication Style')
add_table(['Aspect', 'Detail'], [
    ['Primary method', 'WhatsApp calls + In-person \u2014 "cheaper to use" and "affordable"'],
    ['Who he contacts', 'Neighbours, customers, colleagues, CDA chairman, church members'],
    ['Call duration', 'Very short: "few minutes," "few seconds," "5 minutes," "not long"'],
    ['Conversation style', 'Almost entirely unplanned: "random conversation," "it just came up," "something came up"'],
    ['Cost consciousness', '"Cheaper to use," "affordable" \u2014 WhatsApp calls to save airtime'],
    ['In-person fallback', 'Communicates "in person" when "phone was down" \u2014 forced by dead battery'],
])
doc.add_paragraph(
    'Adekunle\u2019s communication style is reactive and efficient. His conversations are almost '
    'never planned \u2014 they "just come up" as part of his shop life. He favours WhatsApp calls '
    'because they are "cheaper to use," revealing strong cost-consciousness in his communication '
    'choices. His calls are extremely short (seconds to minutes), focused on immediate business '
    'needs. The most telling communication detail is his forced shift to "in person" communication '
    'when his phone died \u2014 proving that without his phone, his social and professional world '
    'shrinks to whoever is physically present.'
)

# ========== DIM 8 ==========
add_h('Dimension 8: Emotional Profile & Stress Map')

add_h3('Daily Emotional Arc')
add_table(['Time Block', 'Emotion', 'Evidence'], [
    ['Wake-up', 'Positive: Happy (3/6), Relaxed (2/6), Tired (1/6)', 'Early riser who starts purposefully; prayer provides stability'],
    ['Morning', 'Focused and hopeful', '"To make sales" \u2014 morning is for preparation and optimism'],
    ['Mid-day', 'Anxious and frustrated', '"Awaiting customers," "low sales" \u2014 uncertainty dominates'],
    ['Afternoon', 'Crisis-driven (on bad days)', '"finding how to charge my phone" \u2014 infrastructure panic'],
    ['Evening', 'Escapist or depleted', 'TikTok when phone works; sleep when it doesn\u2019t'],
])

add_h3('Stress Triggers (Ranked)')
doc.add_paragraph('1. Power outage / Dead phone \u2014 the existential threat to his business operations')
doc.add_paragraph('2. Low sales / No sales \u2014 directly linked to his emotional state')
doc.add_paragraph('3. Government council regulatory pressure \u2014 an external threat he cannot predict')
doc.add_paragraph('4. Rain \u2014 affects his shop and reduces customer traffic')
doc.add_paragraph('5. Hunger / Inability to eat well \u2014 the physical cost of financial scarcity')

add_h3('Resilience Pattern')
doc.add_paragraph(
    'Adekunle\u2019s resilience mechanism is TikTok. On days when his phone is alive, he consistently '
    'turns to TikTok in the evening to "relax and escape stress." This is not idle scrolling \u2014 '
    'it is a deliberate emotional regulation strategy. He also draws strength from his faith: '
    'prayer is his first act every morning, and church attendance provides community. But his '
    'emotional state is fundamentally sales-dependent: "when there\u2019s sales" is his answer to '
    '"what makes you happy" on 5/6 days. Without sales, there is no happiness \u2014 only endurance.'
)

# ========== DIM 9 ==========
add_h('Dimension 9: Opportunities for the Brand')
add_table(['Opportunity', 'Actionable Insight'], [
    ['Ultra-Durable Power Banks', 'Adekunle\u2019s power bank was "faulty" and NEPA provided nothing. He represents an '
     'underserved segment: shop owners with zero grid reliability. A branded, high-capacity power '
     'bank bundled with an MTN recharge deal would solve his #1 problem and create brand lock-in'],
    ['TikTok Data Bundles', 'TikTok is his sole digital escape (3/6 days). A dedicated "TikTok Night Bundle" \u2014 '
     'cheap, high-volume data for evening video streaming \u2014 would match his exact usage pattern'],
    ['Small Business Phone as Multi-Tool', 'He uses his phone as calculator, torch, alarm, and communicator. A rugged, '
     'long-battery smartphone marketed as "the shop owner\u2019s phone" with preloaded calculator, '
     'flashlight, and WhatsApp Business would resonate deeply'],
    ['WhatsApp Business for Traders', 'He communicates with customers via WhatsApp and uses calls for crisis management. '
     'Training programs or simplified WhatsApp Business onboarding for market traders would increase '
     'his efficiency and reduce his reliance on voice calls'],
    ['Regulatory Support Information Line', 'The government council dispute consumed an entire day. An SMS-based information '
     'service for small business permit requirements would help traders like Adekunle navigate '
     'regulatory threats without losing a full day of sales'],
])

# ========== DEFINING QUOTE ==========
doc.add_paragraph()
add_h('Defining Quote', level=2)
add_quote('In the shop with flat battery')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run(
    '\u2014 Five words that summarise the helplessness of a Lagos shopkeeper when infrastructure fails. '
    'Adekunle is physically present at his shop, ready to work, but his phone is dead. Without it, '
    'he cannot calculate prices, cannot receive customer calls, cannot verify payments, cannot even '
    'see in his shop without the torch. He is there but not there. He is open but closed. '
    'The phone battery is not a convenience \u2014 it is the on/off switch for his entire livelihood.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.save(f'{outdir}\\Persona_Adekunle_Adepetun.docx')
print("Persona saved: Adekunle Adepetun (v2 deep-dive)")
